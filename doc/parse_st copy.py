from docx import Document
from pathlib import Path
import re
import json
from zipfile import ZipFile, BadZipFile


# -----------------------------
# Regex patterns
# -----------------------------
JOB_RE = re.compile(r"JOB#\s*([A-Z0-9-]+)", re.I)
PART_RE     = re.compile(r"PART#\s*([A-Z0-9\-]+)", re.I)
REV_RE = re.compile(r"\bREV[:\s]*([A-Z0-9\/]+)", re.I)
PO_RE = re.compile(r"PO#\s*([A-Z0-9\-]+)", re.I)
ORDER_QTY_RE = re.compile(r"ORDER QTY:\s*(\d+)", re.I)
DUE_RE = re.compile(r"DUE DATE:\s*([\d/]+)", re.I)
CUSTOMER_RE = re.compile(r"CUSTOMER:\s*([A-Z0-9]+)", re.I)
RISK_RE = re.compile(r"RISK:\s*(\w+)", re.I)
MAT_RE = re.compile(r"MATERIAL\s*:\s*(.+)", re.I)
PART_NAME_RE = re.compile(r"PART NAME\s*:\s*(.+)", re.I)
NOTE_RE = re.compile(r"NOTE\s*:\s*(.+)", re.I)
REACTION_PLAN_RE = re.compile(r"^REACTION PLAN", re.I)
QTY_SECTION_RE = re.compile(r"^QTY\.?\s+TO\s+(SHIP|STOCK)", re.I)
COMMENTS_RE = re.compile(r"^COMMENTS", re.I)



NUM_STEP_RE = re.compile(r"^\d{3}$")      # 010, 020
MAT_STEP_RE = re.compile(r"^M\d+$")       # M1, M2, M3


# -----------------------------
# Utilities
# -----------------------------
def clean(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def ensure_docx(path: Path):
    try:
        with ZipFile(path):
            pass
    except BadZipFile:
        raise ValueError(
            f"{path} is not a valid .docx file. "
            "Please re-save as Word Document (*.docx)."
        )


def iter_all_text(doc: Document):
    # Normal paragraphs
    for p in doc.paragraphs:
        if p.text.strip():
            yield clean(p.text)

    # Table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        yield clean(p.text)
def strip_md(text: str) -> str:
    return re.sub(r"\*\*(.*?)\*\*", r"\1", text).strip()

def paragraph_to_bold_markdown(p):
    result = ""
    for run in p.runs:
        if not run.text:
            continue

        if run.bold:
            result += f"**{run.text}**"
        else:
            result += run.text

    return result.strip()   # ❌ อย่าเรียก clean()

def iter_all_text_with_bold(doc: Document):
    # paragraphs ปกติ
    for p in doc.paragraphs:
        if p.text.strip():
            yield paragraph_to_bold_markdown(p)

    # ตาราง
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        yield paragraph_to_bold_markdown(p)
def iter_header_text(doc: Document):
    for section in doc.sections:
        header = section.header

        # header paragraphs
        for p in header.paragraphs:
            if p.text.strip():
                yield clean(paragraph_to_bold_markdown(p))

        # header tables
        for table in header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if p.text.strip():
                            yield clean(paragraph_to_bold_markdown(p))
def parse_header_lines(lines, data):
    for line in lines:
        raw = strip_md(line)
       
        if m := JOB_RE.search(raw):
            data["lot"]["lot_no"] = m.group(1)

        elif m := PART_RE.search(raw):
            data["lot"]["part_no"] = m.group(1)

        elif m := REV_RE.search(raw):
            data["lot"]["rev"] = m.group(1)

        elif m := PO_RE.search(raw):
            data["lot"]["po_no"] = m.group(1)

        elif m := ORDER_QTY_RE.search(raw):
            data["lot"]["planned_qty"] = int(m.group(1))

        elif m := DUE_RE.search(raw):
            data["lot"]["due_date"] = m.group(1)

        elif m := CUSTOMER_RE.search(raw):
            
            data["traveler"]["customer_code"] = m.group(1)

       


def extract_material_detail_from_header(doc: Document) -> str | None:
    """
    หา MATERIAL: xxxx จากตาราง header แรก
    """
    for section in doc.sections:
        header = section.header

        for table in header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        text = p.text.strip()
                        if text.upper().startswith("MATERIAL:"):
                            return text.replace("MATERIAL:", "", 1).strip()

    return None

# -----------------------------
# Main parser
# -----------------------------
def parse_docx(path: str) -> dict:
    
    path = Path(path)
   
    if path.suffix.lower() != ".docx":
        raise ValueError("Only .docx files are supported")

    ensure_docx(path)

    doc = Document(path)
    header_lines = list(iter_header_text(doc))
    body_lines = list(iter_all_text_with_bold(doc))

    data = {
        "lot": {},
        "traveler": {},
        "steps": []
    }
   
    # ✅ parse real header ก่อน
    parse_header_lines(header_lines, data)


    

    current_step = None
    step_order = 0

    for line in body_lines:
        # raw_line = strip_md(line)
       
    # 🔑 ต้อง strip markdown ก่อนทุกครั้ง
        raw_line = strip_md(line)

        # -------- LOT / HEADER --------
        if m := JOB_RE.search(raw_line):
            data["lot"]["lot_no"] = m.group(1)
           

        elif m := PART_RE.search(raw_line):
            data["lot"]["part_no"] = m.group(1)
            

        elif m := REV_RE.search(raw_line):
            data["lot"]["rev"] = m.group(1)
            

        elif m := PO_RE.search(raw_line):
            data["lot"]["po_no"] = m.group(1)

        elif m := ORDER_QTY_RE.search(raw_line):
            data["lot"]["planned_qty"] = int(m.group(1))

        elif m := DUE_RE.search(raw_line):
            data["lot"]["due_date"] = m.group(1)

        elif m := CUSTOMER_RE.search(raw_line):
            data["traveler"]["customer_code"] = m.group(1)

        elif m := RISK_RE.search(raw_line):
            data["traveler"]["risk"] = m.group(1)
        elif m := MAT_RE.search(raw_line):
            # print("Found MATERIAL in header:", m.group(1))
            data["lot"]["material_detail"] = m.group(1)
        elif m := PART_NAME_RE.search(raw_line):
            data["lot"]["part_name"] = m.group(1)
        elif m := NOTE_RE.search(raw_line):
            data["lot"]["note"] = m.group(1)
     

        # -------- STEP HEADER --------
        elif NUM_STEP_RE.fullmatch(raw_line) or MAT_STEP_RE.fullmatch(raw_line):

            if current_step:
                data["steps"].append(current_step)

            step_order += 1

            current_step = {
                "order": step_order,
                "step_code": raw_line,   # ✅ ใช้ raw_line
                "step_type": "material" if raw_line.startswith("M") else "process",
                "step_name": "",
                "notes": "",
                "qa_required": False
            }

        # -------- STEP BODY --------
        elif current_step:

            # ❌ ถ้าเข้า reaction plan = จบ step
            if (
                REACTION_PLAN_RE.match(raw_line)
                or QTY_SECTION_RE.match(raw_line)
                or COMMENTS_RE.match(raw_line)
            ):
                data["steps"].append(current_step)
                current_step = None
                continue

            if not current_step["step_name"]:
                current_step["step_name"] = line
            else:
                current_step["notes"] += line + "\n"


    if current_step:
        data["steps"].append(current_step)

    # -------- TRAVELER --------
    data["traveler"]["traveler_no"] = data["lot"].get("lot_no")
    material_detail = extract_material_detail_from_header(doc)
    if material_detail:
        data["lot"]["material_detail"] = material_detail

    return data


# -----------------------------
# CLI runner
# -----------------------------
if __name__ == "__main__":
    import json
    from pathlib import Path

    path = Path(r"C:\Users\Tanapon\Documents\GitHub\cnc\doc")

    template_path = path / "L16492.docx"
    

    result = parse_docx(template_path)
    part_no = result["lot"].get("part_no","UNKNOWN")
    part_rev = result["lot"].get("rev","-")
    lot_no = result["lot"].get("lot_no","UNKNOWN")  
    
    output_json = path / f"{part_no}_{part_rev}_{lot_no}.json"

    with open(output_json, "w", encoding="utf-8") as f:
        json.dump(result, f, indent=2, ensure_ascii=False)

    print(f"✅ Saved parsed data to: {output_json}")

