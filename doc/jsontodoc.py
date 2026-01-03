from docx import Document
from pathlib import Path
import json
import copy
import re


# --------------------------------------------------
# Helper: clone a Word table row (safe)
# --------------------------------------------------
def clone_row(table, row_idx):
    row = table.rows[row_idx]
    new_tr = copy.deepcopy(row._tr)

    pos = table._tbl.index(row._tr)
    table._tbl.insert(pos + 1, new_tr)
    return table.rows[row_idx + 1]

# def print_row(row, row_idx=None):
#     if row_idx is not None:
#         print(f"Row {row_idx}:")
#     else:
#         print("Row:")

#     for i, cell in enumerate(row.cells):
#         text = cell.text.replace("\n", "\\n")
#         print(f"  Cell {i}: {text}")   

# def print_table_row(table, row_idx):
    
#     row = table.rows[row_idx]
#     print(f"Row {row_idx}:")
#     for i, cell in enumerate(row.cells):
#         if i>1:
#             break   
#         text = cell.text.replace("\n", "\\n")
#         print(f"  Cell {i}: {text}")
# --------------------------------------------------
# Helper: find row containing specific text
# --------------------------------------------------
def find_row_index(table, keyword):
    for i, row in enumerate(table.rows):
        text = "\n".join(cell.text for cell in row.cells)
        if keyword in text:
            return i
    return None


def find_row_object(table, keyword):
    for row in table.rows:
        text = "\n".join(cell.text for cell in row.cells)
        if keyword in text:
            return row
    return None

def clone_row_at(table, src_row_idx, insert_at):
    src_row = table.rows[src_row_idx]
    new_row = copy.deepcopy(src_row._tr)
    table._tbl.insert(insert_at, new_row)
    return table.rows[insert_at]
def replace_placeholder_fuzzy(paragraph, placeholder, value):
    if not paragraph.runs:
        return

    full_text = "".join(run.text for run in paragraph.runs)

    pattern = re.escape(placeholder[0])
    for ch in placeholder[1:]:
        pattern += r"\s*" + re.escape(ch)

    if not re.search(pattern, full_text):
        return

    new_text = re.sub(pattern, str(value), full_text)

    for run in paragraph.runs:
        run.text = ""

    paragraph.runs[0].text = new_text
def replace_header(doc, mapping):
    for section in doc.sections:
        header = section.header

        for p in header.paragraphs:
            for k, v in mapping.items():
                replace_placeholder_fuzzy(p, k, v)

        for table in header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for k, v in mapping.items():
                            replace_placeholder_fuzzy(p, k, v)

def replace_body(doc, mapping):
    # print(mapping)
    # ---------- paragraphs (body text) ----------
    for p in doc.paragraphs:
        for key, value in mapping.items():
            if key in p.text:
                for run in p.runs:
                    run.text = run.text.replace(key, str(value))

    # ---------- tables (body tables) ----------
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for key, value in mapping.items():
                        if key in p.text:
                            for run in p.runs:
                                run.text = run.text.replace(key, str(value))
def replace_body_fuzzy(doc, mapping):
    for p in doc.paragraphs:
        for k, v in mapping.items():
            replace_placeholder_fuzzy(p, k, v)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for k, v in mapping.items():
                        replace_placeholder_fuzzy(p, k, v)
# --------------------------------------------------
# MAIN
# --------------------------------------------------
def generate_traveler(template_path, json_path, output_path):

    template_path = Path(template_path)
    json_path = Path(json_path)
    output_path = Path(output_path).resolve()

    output_path.parent.mkdir(parents=True, exist_ok=True)
    if output_path.exists():
        output_path.unlink()

    doc = Document(str(template_path))
    data = json.loads(json_path.read_text(encoding="utf-8"))

    # -------------------------
    # HEADER
    # -------------------------
    header_map = {
        "{{part}}": data["lot"]["part_no"],
        "{{lot}}": data["lot"]["lot_no"],
        "{{po}}": data["lot"]["po_no"],
        "{{due}}": data["lot"]["due_date"],
        "{{cus}}": data["traveler"]["customer_code"],
        "{{qty}}": data["lot"]["planned_qty"],
        "{{material_detail}}": data["lot"]["material_detail"],
    }
    replace_header(doc, header_map)
    # replace_body(doc, header_map) 
    replace_body_fuzzy(doc, header_map)
    # -------------------------
    # MATERIAL (M1 / M2)
    # -------------------------
    for step in data["steps"]:
        if step["step_code"] in ("M1", "M2"):
            prefix = step["step_code"]

            mat_map = {
                f"{{{{{prefix}_MATERIAL_SIZE}}}}": step.get("material_size", ""),
                f"{{{{{prefix}_TOTAL_LENGTH}}}}": step.get("total_length", ""),
                f"{{{{{prefix}_SUPPLIER}}}}": step.get("supplier", ""),
                f"{{{{{prefix}_PO}}}}": step.get("po", ""),
                f"{{{{{prefix}_HT}}}}": step.get("ht", ""),
            }

            # 🔴 ใช้ fuzzy replace เท่านั้น
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for k, v in mat_map.items():
                                replace_placeholder_fuzzy(p, k, v)




    # --------------------------------------------------
    # Find STEP TABLE, M2 row, and {op} template row
    # --------------------------------------------------
    step_table = None
    m2_row_idx = None
    
    for table in doc.tables:
        for i, row in enumerate(table.rows):
            text = "\n".join(cell.text for cell in row.cells)

            if text.strip().startswith("M2"):
                m2_row_idx = i
                print("Found M2 row at", i)

            if "{op}" in text:
                step_table = table
                print("Found OP template row")

        if step_table and m2_row_idx is not None:
            break

    if step_table is None or m2_row_idx is None:
        raise RuntimeError("Cannot find M2 row or {op} template row")
    

    # --------------------------------------------------
    # Insert PROCESS steps AFTER M2
    # --------------------------------------------------
    insert_at = m2_row_idx  + 1
    
    # print("Inserting process steps after row", insert_at, "which is M2", m2_row_idx   )
    for step in data["steps"]:
       
        if step.get("step_type") != "process":
            continue

        # print(f"Inserting step {step['step_code']} after M2 at row {insert_at}" , m2_row_idx)
        new_row = clone_row(step_table, insert_at )
        
        # Column 0 = OP code
        new_row.cells[0].text = step["step_code"]

        cell = new_row.cells[1]

        # ล้างข้อความ template เดิม
        cell.text = ""

        # ---------- STEP NAME (markdown bold) ----------
        p_name = cell.paragraphs[0]
        add_text_with_bold(p_name, step["step_name"])

        # # ---------- NOTES (markdown bold) ----------
        # notes = step.get("notes", "")
        # if notes:
        #     p_notes = cell.add_paragraph()
        #     add_text_with_bold(p_notes, notes)


        
        notes = step.get("notes", "")
        p = new_row.cells[1].add_paragraph()
        add_text_with_bold(p, step.get("notes", ""))
        # if notes:
        #     new_row.cells[1].add_paragraph(notes)
        
        insert_at += 1
    
    # --------------------------------------------------
    # Remove {op} template row SAFELY (no index bug)
    # --------------------------------------------------
    step_table = None
    remove_idx = None
    for table in doc.tables:
        for i, row in enumerate(table.rows):
            text = "\n".join(cell.text for cell in row.cells)
            if "{op}" in text:
                step_table = table
                remove_idx = i
                print("Found OP template row")

        if step_table  is not None:
            break 
    table._tbl.remove(step_table.rows[remove_idx]._tr)  

    # --------------------------------------------------
    # Save output
    # --------------------------------------------------
    doc.save(str(output_path))
    print(f"Shop Traveler created successfully: {output_path}")

import re

def add_text_with_bold(paragraph, text):
    """
    รองรับ **bold**
    """
    parts = re.split(r"(\*\*.*?\*\*)", text)

    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            r.bold = True
        else:
            paragraph.add_run(part)
# --------------------------------------------------
# CLI
# --------------------------------------------------
if __name__ == "__main__":

    # template_path = r"C:\Users\TPSERVER\dev\cnc_project\doc\shop_templete.docx"
    # json_path = r"C:\Users\TPSERVER\dev\cnc_project\doc\data.json"
    # output_path = r"C:\Users\TPSERVER\dev\cnc_project\doc\output_traveler.docx"

    path = r"C:\Users\Tanapon\Documents\GitHub\cnc\doc"
    template_path = path + r"\shop_templete.docx"
    json_path = path + r"\data.json"
    output_path = path + r"\output_traveler.docx"

    generate_traveler(template_path, json_path, output_path)
