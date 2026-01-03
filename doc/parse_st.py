from docx import Document
from pathlib import Path
import re
import json
from zipfile import ZipFile, BadZipFile


# -----------------------------
# Regex patterns
# -----------------------------
JOB_RE = re.compile(r"JOB#\s*([A-Z0-9-]+)", re.I)
PART_RE = re.compile(r"PART\s*#?:\s*([A-Z0-9\-]+)", re.I)
REV_RE = re.compile(r"\bREV[:\s]*([A-Z0-9\/]+)", re.I)
PO_RE = re.compile(r"PO#\s*([A-Z0-9\-]+)", re.I)
ORDER_QTY_RE = re.compile(r"ORDER QTY:\s*(\d+)", re.I)
DUE_RE = re.compile(r"DUE DATE:\s*([\d/]+)", re.I)
CUSTOMER_RE = re.compile(r"CUSTOMER:\s*([A-Z0-9]+)", re.I)
RISK_RE = re.compile(r"RISK:\s*(\w+)", re.I)

NUM_STEP_RE = re.compile(r"^\d{3}$")      # 010, 020
MAT_STEP_RE = re.compile(r"^M\d+$")       # M1, M2, M3


# -----------------------------
# Utilities
# -----------------------------
def clean(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def ensure_docx(path: Path):
    try:
        with ZipFile(path):
            pass
    except BadZipFile:
        raise ValueError(
            f"{path} is not a valid .docx file. "
            "Please re-save as Word Document (*.docx)."
        )


def iter_all_text(doc: Document):
    # Normal paragraphs
    for p in doc.paragraphs:
        if p.text.strip():
            yield clean(p.text)

    # Table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        yield clean(p.text)


# -----------------------------
# Main parser
# -----------------------------
def parse_docx(path: str) -> dict:
    path = Path(path)

    if path.suffix.lower() != ".docx":
        raise ValueError("Only .docx files are supported")

    ensure_docx(path)

    doc = Document(path)
    lines = list(iter_all_text(doc))

    data = {
        "lot": {},
        "traveler": {},
        "steps": []
    }

    current_step = None
    step_order = 0

    for line in lines:

        # -------- LOT / HEADER --------
        if m := JOB_RE.search(line):
            data["lot"]["lot_no"] = m.group(1)

        elif m := PART_RE.search(line):
            data["lot"]["part_no"] = m.group(1)

        elif m := REV_RE.search(line):
            data["lot"]["rev"] = m.group(1)

        elif m := PO_RE.search(line):
            data["lot"]["po_no"] = m.group(1)

        elif m := ORDER_QTY_RE.search(line):
            data["lot"]["planned_qty"] = int(m.group(1))

        elif m := DUE_RE.search(line):
            data["lot"]["due_date"] = m.group(1)

        elif m := CUSTOMER_RE.search(line):
            data["traveler"]["customer_code"] = m.group(1)

        elif m := RISK_RE.search(line):
            data["traveler"]["risk"] = m.group(1)

        # -------- STEP HEADER --------
        elif NUM_STEP_RE.fullmatch(line) or MAT_STEP_RE.fullmatch(line):

            if current_step:
                data["steps"].append(current_step)

            step_order += 1

            current_step = {
                "order": step_order,
                "step_code": line,
                "step_type": "material" if line.startswith("M") else "process",
                "step_name": "",
                "notes": "",
                "qa_required": False
            }

        # -------- STEP BODY --------
        elif current_step:
            if not current_step["step_name"]:
                current_step["step_name"] = line
            else:
                current_step["notes"] += line + "\n"

            if "QA TO" in line.upper():
                current_step["qa_required"] = True

    if current_step:
        data["steps"].append(current_step)

    # -------- TRAVELER --------
    data["traveler"]["traveler_no"] = data["lot"].get("lot_no")

    return data


# -----------------------------
# CLI runner
# -----------------------------
if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Usage: python parse_st.py <file.docx>")
        sys.exit(1)

    result = parse_docx(sys.argv[1])
    print(json.dumps(result, indent=2))
