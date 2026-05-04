# services/traveler_docx.py
from docx import Document
import copy
import re
from pathlib import Path

# ==================================================
# Helpers
# ==================================================
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

def center_cell(cell):
    # แนวตั้ง (vertical center)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    # แนวนอน (horizontal center)
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
def clone_row(table, row_idx):
    row = table.rows[row_idx]
    new_row = copy.deepcopy(row._tr)
    table._tbl.insert(row_idx + 1, new_row)
    return table.rows[row_idx + 1]


def add_text_with_bold(paragraph, text):
    """
    รองรับ **bold**
    """
    parts = re.split(r"(\*\*.*?\*\*)", text or "")
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            r.bold = True
        else:
            paragraph.add_run(part)


def replace_placeholder_fuzzy(paragraph, placeholder, value):
    if not paragraph.runs:
        return

    full_text = "".join(run.text for run in paragraph.runs)

    pattern = re.escape(placeholder[0])
    for ch in placeholder[1:]:
        pattern += r"\s*" + re.escape(ch)

    if not re.search(pattern, full_text):
        return

    new_text = re.sub(pattern, str(value), full_text)

    for run in paragraph.runs:
        run.text = ""

    paragraph.runs[0].text = new_text


def replace_header(doc, mapping):
    for section in doc.sections:
        header = section.header

        # header paragraphs
        for p in header.paragraphs:
            for k, v in mapping.items():
                replace_placeholder_fuzzy(p, k, v)

        # header tables
        for table in header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for k, v in mapping.items():
                            replace_placeholder_fuzzy(p, k, v)


def replace_body_fuzzy(doc, mapping):
    # body paragraphs
    for p in doc.paragraphs:
        for k, v in mapping.items():
            replace_placeholder_fuzzy(p, k, v)

    # body tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for k, v in mapping.items():
                        replace_placeholder_fuzzy(p, k, v)

def safe_text(v):
    if v is None:
        return ""
    return str(v)
# ==================================================
# MAIN GENERATOR (DB-BASED)
# ==================================================


from docx import Document
from pathlib import Path
import qrcode
from docx.shared import Inches

def generate_qr(data, path):
    img = qrcode.make(data)
    img.save(path)


def replace_qr(doc, placeholder, image_path):
    for p in doc.paragraphs:
        if placeholder in p.text:
            p.text = p.text.replace(placeholder, "")
            run = p.add_run()
            run.add_picture(image_path, width=Inches(1))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if placeholder in p.text:
                        p.text = p.text.replace(placeholder, "")
                        run = p.add_run()
                        run.add_picture(image_path, width=Inches(1))

def generate_traveler_from_db(template_path, data: dict, output_path):
    from docx import Document
    from pathlib import Path

    template_path = Path(template_path)
    output_path = Path(output_path)

    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document(str(template_path))

    # --------------------------------------------------
    # HEADER
    # --------------------------------------------------
    header_map = {
        "{{part}}": data["header"]["part_no"],
        "{{part_name}}": data["header"]["part_name"],
        "{{rev}}": data["header"]["part_rev"],
        "{{lot}}": data["header"]["lot_no"],
        "{{po}}": data["header"]["po_no"],
        "{{cus}}": data["header"]["customer_code"],
        "{{due}}": data["header"]["due_date"],
        "{{qty}}": data["header"]["planned_qty"],
        "{{release}}": data["header"]["release_date"],
        "{{material_detail}}": data["header"]["material_detail"],
    }

    replace_header(doc, header_map)
    replace_body_fuzzy(doc, header_map)

    # -------------------------
    # QR CODE
    # -------------------------
    from tempfile import NamedTemporaryFile

    qr_text = data["header"]["lot_no"]

    tmp_qr = NamedTemporaryFile(suffix=".png", delete=False)
    qr_path = tmp_qr.name
    tmp_qr.close()

    generate_qr(qr_text, qr_path)

    replace_qr(doc, "{{QR}}", qr_path)

    # --------------------------------------------------
    # FIND STEP TABLE + TEMPLATE MARKER
    # --------------------------------------------------
    step_table = None
    template_row_idx = None

    for table in doc.tables:
        for i, row in enumerate(table.rows):
            text = "\n".join(cell.text for cell in row.cells)

            if "{STEP_TEMPLATE}" in text:
                step_table = table
                template_row_idx = i
                break

        if step_table:
            break

    if step_table is None or template_row_idx is None:
        raise RuntimeError("❌ Cannot find {STEP_TEMPLATE} in Word template")

    # --------------------------------------------------
    # START INSERT AFTER MARKER
    # --------------------------------------------------
    insert_at = template_row_idx + 1

    # --------------------------------------------------
    # SORT STEPS (🔥 FIX M1 missing)
    # --------------------------------------------------
    steps_sorted = sorted(data["steps"], key=lambda x: x.get("seq", 0))

    # print("DEBUG steps:", steps_sorted)

    # --------------------------------------------------
    # INSERT STEPS
    # --------------------------------------------------
    for step in steps_sorted:
        # print(f"Inserting: {step.get('step_code')}")

        new_row = clone_row(step_table, insert_at)

        # -------------------------
        # OP CODE
        # -------------------------
        new_row.cells[0].text = str(step.get("step_code", ""))  #  op col 1

        # -------------------------
        # DESCRIPTION (bold support)
        # -------------------------
        desc_cell = new_row.cells[1]        ## Step description col 2
        desc_cell.text = ""

        p = desc_cell.paragraphs[0]

        text = step.get("step_name", "") or ""

        if step.get("step_detail"):
            text += "\n" + step["step_detail"]

        add_text_with_bold(p, text)

        # -------------------------
        # RECEIVE / ACCEPT / REJECT
        # -------------------------
        new_row.cells[2].text = str(step.get("qty_receive", ""))    # receive col 3

        new_row.cells[3].text = str(step.get("operator", ""))   # accept col 6
        new_row.cells[4].text = "\nSUPPLIER: _________\nPO#: _________\nHT#: _________" if "M1" in step.get("step_code", "") or "M2" in step.get("step_code", "") else ""   # Description


        new_row.cells[5].text = str(step.get("qty_accept", ""))   # accept col 6
        new_row.cells[6].text = str(step.get("qty_reject", ""))   # reject col 7

        new_row.cells[7].text = str(step.get("created_at", ""))   # date col 8

        # -------------------------
        # QA (optional)
        # -------------------------
        # if len(new_row.cells) > 5:
        #     new_row.cells[5].text = str(step.get("qa_required", ""))

        # -------------------------
        # CENTER ALIGN
        # -------------------------
        for idx in [0, 2, 3, 5,6, 7]:  # center align for op, receive, accept, reject, qa
            if idx < len(new_row.cells):
                center_cell(new_row.cells[idx])

        insert_at += 1

    # --------------------------------------------------
    # REMOVE TEMPLATE ROWS (🔥 FIX {STEP_TEMPLATE} issue)
    # --------------------------------------------------

    # remove marker row
    for row in step_table.rows:
        text = "\n".join(cell.text for cell in row.cells)
        if "{STEP_TEMPLATE}" in text:
            step_table._tbl.remove(row._tr)
            break

    # remove template row ({{op}})
    for row in step_table.rows:
        text = "\n".join(cell.text for cell in row.cells)
        if "{{op}}" in text:
            step_table._tbl.remove(row._tr)
            break

    # ✅ remove row 2 (index = 1)
    if len(step_table.rows) > 1:
        step_table._tbl.remove(step_table.rows[1]._tr)

    # --------------------------------------------------
    # SAVE
    # --------------------------------------------------
    doc.save(str(output_path))



def generate_traveler_from_db_blank(template_path, data: dict, output_path):
    from docx import Document
    from pathlib import Path

    template_path = Path(template_path)
    output_path = Path(output_path)

    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document(str(template_path))

    # --------------------------------------------------
    # HEADER
    # --------------------------------------------------
    header_map = {
        "{{part}}": data["header"]["part_no"],
        "{{part_name}}": data["header"]["part_name"],
        "{{rev}}": data["header"]["part_rev"],
        "{{lot}}": data["header"]["lot_no"],
        "{{po}}": data["header"]["po_no"],
        "{{cus}}": data["header"]["customer_code"],
        "{{due}}": data["header"]["due_date"],
        "{{qty}}": data["header"]["planned_qty"],
        "{{release}}": data["header"]["release_date"],
        "{{material_detail}}": data["header"]["material_detail"],
    }

    replace_header(doc, header_map)
    replace_body_fuzzy(doc, header_map)

    # -------------------------
    # QR CODE
    # -------------------------
    from tempfile import NamedTemporaryFile

    qr_text = data["header"]["lot_no"]

    tmp_qr = NamedTemporaryFile(suffix=".png", delete=False)
    qr_path = tmp_qr.name
    tmp_qr.close()

    generate_qr(qr_text, qr_path)

    replace_qr(doc, "{{QR}}", qr_path)

    # --------------------------------------------------
    # FIND STEP TABLE + TEMPLATE MARKER
    # --------------------------------------------------
    step_table = None
    template_row_idx = None

    for table in doc.tables:
        for i, row in enumerate(table.rows):
            text = "\n".join(cell.text for cell in row.cells)

            if "{STEP_TEMPLATE}" in text:
                step_table = table
                template_row_idx = i
                break

        if step_table:
            break

    if step_table is None or template_row_idx is None:
        raise RuntimeError("❌ Cannot find {STEP_TEMPLATE} in Word template")

    # --------------------------------------------------
    # START INSERT AFTER MARKER
    # --------------------------------------------------
    insert_at = template_row_idx + 1

    # --------------------------------------------------
    # SORT STEPS (🔥 FIX M1 missing)
    # --------------------------------------------------
    steps_sorted = sorted(data["steps"], key=lambda x: x.get("seq", 0))

    # print("DEBUG steps:", steps_sorted)

    # --------------------------------------------------
    # INSERT STEPS
    # --------------------------------------------------
    for step in steps_sorted:
        # print(f"Inserting: {step.get('step_code')}")

        new_row = clone_row(step_table, insert_at)

        # -------------------------
        # OP CODE
        # -------------------------
        new_row.cells[0].text = str(step.get("step_code", ""))  #  op col 1

        # -------------------------
        # DESCRIPTION (bold support)
        # -------------------------
        desc_cell = new_row.cells[1]        ## Step description col 2
        desc_cell.text = ""

        p = desc_cell.paragraphs[0]

        text = step.get("step_name", "") or ""

        if step.get("step_detail"):
            text += "\n" + step["step_detail"]

        add_text_with_bold(p, text)

        # -------------------------
        # RECEIVE / ACCEPT / REJECT
        # -------------------------
        # new_row.cells[2].text = str(step.get("qty_receive", ""))    # receive col 3

        # new_row.cells[3].text = str(step.get("operator", ""))   # accept col 6
        new_row.cells[4].text = "\nSUPPLIER: _________\nPO#: _________\nHT#: _________" if "M1" in step.get("step_code", "") or "M2" in step.get("step_code", "") else ""   # Description


        # new_row.cells[5].text = str(step.get("qty_accept", ""))   # accept col 6
        # new_row.cells[6].text = str(step.get("qty_reject", ""))   # reject col 7

        # new_row.cells[7].text = str(step.get("created_at", ""))   # date col 8

        # -------------------------
        # QA (optional)
        # -------------------------
        # if len(new_row.cells) > 5:
        #     new_row.cells[5].text = str(step.get("qa_required", ""))

        # -------------------------
        # CENTER ALIGN
        # -------------------------
        for idx in [0, 2, 3, 5,6, 7]:  # center align for op, receive, accept, reject, qa
            if idx < len(new_row.cells):
                center_cell(new_row.cells[idx])

        insert_at += 1

    # --------------------------------------------------
    # REMOVE TEMPLATE ROWS (🔥 FIX {STEP_TEMPLATE} issue)
    # --------------------------------------------------

    # remove marker row
    for row in step_table.rows:
        text = "\n".join(cell.text for cell in row.cells)
        if "{STEP_TEMPLATE}" in text:
            step_table._tbl.remove(row._tr)
            break

    # remove template row ({{op}})
    for row in step_table.rows:
        text = "\n".join(cell.text for cell in row.cells)
        if "{{op}}" in text:
            step_table._tbl.remove(row._tr)
            break

    # ✅ remove row 2 (index = 1)
    if len(step_table.rows) > 1:
        step_table._tbl.remove(step_table.rows[1]._tr)

    # --------------------------------------------------
    # SAVE
    # --------------------------------------------------
    doc.save(str(output_path))

def generate_inspection_from_db(template_path, data: dict, output_path):
    from docx import Document
    from pathlib import Path
    from collections import defaultdict
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

    template_path = Path(template_path)
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    print("Inspection data:", data)

    # =========================
    # SAFE INT
    # =========================
    def safe_int(v):
        try:
            return int(v)
        except:
            return 9999

    # =========================
    # GROUP BY OP
    # =========================
    grouped = defaultdict(list)

    for item in data["items"]:
        if item.get("bb_no") in ["Bubble #", None]:
            continue
        grouped[item["op_no"]].append(item)

    # =========================
    # SORT OP + BUBBLE
    # =========================
    ops = sorted(grouped.keys(), key=lambda x: safe_int(x))

    for op in grouped:
        grouped[op].sort(key=lambda x: safe_int(x.get("bb_no")))

    # =========================
    # SPLIT INTO PAGES (3 OP)
    # =========================
    def chunk(lst, n=3):
        for i in range(0, len(lst), n):
            yield lst[i:i+n]

    pages = list(chunk(ops, 3))

    # =========================
    # TEMPLATE CONFIG
    # =========================
    MAX_DATA_ROWS = 13   # row 14 = Note

    # =========================
    # HELPERS
    # =========================
    def find_header_row(table):
        for i, row in enumerate(table.rows):
            texts = [c.text.strip() for c in row.cells]
            if any("B/B" in t for t in texts):
                return i
        return 2

    def clear_data_rows(table, start_row):
        for r in range(start_row, len(table.rows)):
            for c in range(len(table.columns)):
                table.cell(r, c).text = ""

    def center(cell):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # =========================
    # HEADER MAP
    # =========================
    header_map = {
        "{{part}}": data["lot"].get("part_no", ""),
        "{{rev}}": data["lot"].get("part_rev", ""),
        "{{lot}}": data["lot"].get("lot_no", ""),
        "{{po}}": data["lot"].get("po_no", ""),
    }

    # =========================
    # MAIN LOOP
    # =========================
    final_doc = None

    for page_idx, page_ops in enumerate(pages):

        # 🔥 IMPORTANT: new doc per page
        doc = Document(str(template_path))
        replace_header(doc, header_map)

        table = doc.tables[0]

        header_row = find_header_row(table)
        DATA_START_ROW = header_row + 1
        NOTE_ROW_IDX = DATA_START_ROW + MAX_DATA_ROWS

        clear_data_rows(table, DATA_START_ROW)

        # =========================
        # OP HEADER
        # =========================
        op_header_row = header_row - 1

        for col_idx, op in enumerate(page_ops):
            base_col = col_idx * 4

            table.cell(op_header_row, base_col).text = "date"
            table.cell(op_header_row, base_col + 1).text = op
            table.cell(op_header_row, base_col + 2).text = "emp"

            

        # =========================
        # LIMIT ROWS (avoid NOTE overwrite)
        # =========================
        max_rows = min(
            max(len(grouped[op]) for op in page_ops),
            MAX_DATA_ROWS
        )

        # =========================
        # FILL DATA
        # =========================
        for row_idx in range(max_rows):
            for col_idx, op in enumerate(page_ops):

                base_col = col_idx * 4
                items = grouped[op]

                if row_idx < len(items):
                    item = items[row_idx]
                    r = DATA_START_ROW + row_idx

                    table.cell(r, base_col + 0).text = str(item.get("bb_no", "") or "")
                    table.cell(r, base_col + 1).text = str(item.get("dimension", "") or "")
                    table.cell(r, base_col + 2).text = str(item.get("tqw", "") or "")
                    table.cell(r, base_col + 3).text = "✓" if item.get("fa") else ""

                    for i in range(4):
                        center(table.cell(r, base_col + i))

        # =========================
        # NOTE ROW (PER OP)
        # =========================
        for col_idx, op in enumerate(page_ops):

            base_col = col_idx * 4

            # ✅ Column 1 → "Note"
            note_cell = table.cell(NOTE_ROW_IDX, base_col)
            note_cell.text = "Note"

            for p in note_cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # ✅ Columns 2–4 → merge into one
            merge_cell = table.cell(NOTE_ROW_IDX, base_col + 1)

            for i in range(2, 4):
                merge_cell = merge_cell.merge(
                    table.cell(NOTE_ROW_IDX, base_col + i)
                )

            merge_cell.text = ""  # empty space for writing

            for p in merge_cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # =========================
        # MERGE DOCS
        # =========================
        if page_idx == 0:
            final_doc = doc
        else:
            final_doc.add_page_break()
            for el in doc.element.body:
                final_doc.element.body.append(el)

    # =========================
    # SAVE
    # =========================
    final_doc.save(str(output_path))