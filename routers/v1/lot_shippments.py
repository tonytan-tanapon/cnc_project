from fastapi import APIRouter, Depends, HTTPException, Body
from sqlalchemy.orm import Session, joinedload

from sqlalchemy import func
from datetime import datetime
from pydantic import BaseModel

from database import get_db
from models import (
    POLine,
    ProductionLot,
    CustomerShipment,
    CustomerShipmentItem,
    Part,
    PartRevision,PO
)

router = APIRouter(prefix="/lot-shippments", tags=["lot-shippments"])

class AllocatePartRequest(BaseModel):
    source_lot_id: int
    target_lot_id: int
    qty: float
    shipment_id: int | None = None  # ✅ ใหม่
# ---------- Helper ----------
def get_lot_or_404(db: Session, lot_id: int) -> ProductionLot:
    lot = db.get(ProductionLot, lot_id)
    if not lot:
        raise HTTPException(status_code=404, detail="Lot not found")
    return lot


def get_part_inventory_data(db: Session, lot_id: int):
    lot = get_lot_or_404(db, lot_id)
    part = db.get(Part, lot.part_id)
    if not part:
        raise HTTPException(status_code=404, detail="Part not found")

    po = db.get(PO, lot.po_id) if lot.po_id else None
    po_number = po.po_number if po else None

    from models import ShopTraveler, ShopTravelerStep

    # ===== คิด finished_qty จาก traveler steps =====
    sub_last_step = (
        db.query(
            ShopTravelerStep.traveler_id,
            func.max(ShopTravelerStep.seq).label("max_seq"),
        )
        .group_by(ShopTravelerStep.traveler_id)
        .subquery()
    )

    finished_qty = (
        db.query(func.coalesce(func.sum(ShopTravelerStep.qty_accept), 0))
        .join(
            sub_last_step,
            (ShopTravelerStep.traveler_id == sub_last_step.c.traveler_id) &
            (ShopTravelerStep.seq == sub_last_step.c.max_seq)
        )
        .join(ShopTraveler, ShopTraveler.id == ShopTravelerStep.traveler_id)
        .filter(ShopTraveler.lot_id == lot.id)
        .filter(ShopTravelerStep.status == "passed")
        .scalar()
        or 0
    )

    # ===== shipped_qty จาก lot_allocate_id =====
    shipped_qty = (
        db.query(func.coalesce(func.sum(CustomerShipmentItem.qty), 0))
        .filter(CustomerShipmentItem.lot_allocate_id == lot_id)
        .scalar()
        or 0
    )

    planned_qty = float(lot.planned_qty or 0)
    finished_qty = float(finished_qty)
    shipped_qty = float(shipped_qty)

    available_qty = max(finished_qty - shipped_qty, 0)

    return {
        "part": part,
        "po_number": po_number,   # ✅ added
        "planned_qty": planned_qty,
        "finished_qty": finished_qty,
        "shipped_qty": shipped_qty,
        "available_qty": available_qty,
    }


# ============================================================
# 1️⃣  List shipments for a lot
# ============================================================
from models import PO, Customer
from models import CustomerShipmentItem  # อย่าลืม import

@router.get("/{lot_id}")
def list_lot_shipments(lot_id: int, db: Session = Depends(get_db)):
    shipments = (
        db.query(CustomerShipment)
        .options(
            joinedload(CustomerShipment.items).joinedload(CustomerShipmentItem.allocated_lot),
            # joinedload(CustomerShipment.items).joinedload(CustomerShipmentItem.lot),  # preload source lot ด้วย
            joinedload(CustomerShipment.items).joinedload(CustomerShipmentItem.lot).joinedload(ProductionLot.part_revision),  # 👈 preload rev
            joinedload(CustomerShipment.po).joinedload(PO.customer),
            joinedload(CustomerShipment.items)
            .joinedload(CustomerShipmentItem.po_line)
            .joinedload(POLine.rev),
        )
        .filter(CustomerShipment.lot_id == lot_id)
        .order_by(CustomerShipment.shipped_at.desc())
        .all()
    )

    result = []
    for s in shipments:
        items = s.items or []
        qty_sum = sum(float(i.qty or 0) for i in items)

        # ✅ GROUP by source lot_no
        lot_group = {}
        for i in items:
            if i.lot and i.lot_id:
                lot_no = i.lot.lot_no
                if lot_no not in lot_group:
                    lot_group[lot_no] = {
                        "qty": 0,
                        "lot_id": i.lot.id,
                        "lot_allocate_id": i.lot_allocate_id  # 👈 ดึง lot_allocate_id เผื่อ UI ต้องใช้ด้วย
                    }
                lot_group[lot_no]["qty"] += float(i.qty or 0)

        # ✅ GROUP by allocated lot to include lot_allocate_id
        allocated_group = {}
        for i in items:
            if i.allocated_lot and i.lot_allocate_id:
                a_lot = i.allocated_lot
                key = a_lot.lot_no
                if key not in allocated_group:
                    allocated_group[key] = {
                        "qty": 0,
                        "lot_id": a_lot.id,
                        "lot_allocate_id": i.lot_allocate_id  # 👈 ฟิลด์ที่ต้องการจาก table
                    }
                allocated_group[key]["qty"] += float(i.qty or 0)

        # ✅ list source lots (ปรับ JSON ให้มี key ตรง)
        lots_list = [
            {
                "lot_no": k,
                "qty": v["qty"],
                "lot_id": v["lot_id"],
                "lot_allocate_id": v["lot_allocate_id"],  # อาจจะ null ได้
            }
            for k, v in lot_group.items()
        ]

        # ✅ list allocated lots พร้อม lot_allocate_id
        allocated_lots_list = [
            {
                "lot_no": k,
                "qty": v["qty"],
                "lot_id": v["lot_id"],
                "lot_allocate_id": v["lot_allocate_id"],  # 👈 ค่า allocate จริง
            }
            for k, v in allocated_group.items()
        ]

        rev = None
        for i in items:
            if i.po_line and i.po_line.rev:
                rev = i.po_line.rev.rev
                break

        rev_po = list({
            i.po_line.rev.rev
            for i in items
            if i.po_line and i.po_line.rev
        })

        rev_lot = list({
            i.lot.part_revision.rev
            for i in items
            if i.lot and i.lot.part_revision
        })

        print(f"Shipment {s.id} has rev from PO lines: {rev_po}, rev from lot revisions: {rev_lot}")

        result.append({
            "id": s.id,
            "shipment_no": s.package_no or f"SHP-{s.id:05d}",
            "qty": qty_sum,
            "uom": "pcs",
            "status": s.status or "pending",
            "date": s.shipped_at,
            "tracking_number": s.tracking_no,
            "customer_name": s.po.customer.name if s.po and s.po.customer else None,
            "customer_code": s.po.customer.code if s.po and s.po.customer else None,

            "lots": lots_list,
            "allocated_lots": allocated_lots_list,

            # ✅ FIXED: rev from PO line
            "rev": rev_po,
            "rev_from_po": rev_po,
            "rev_from_lot": rev_lot,
        })

    return result



# ============================================================
# 2️⃣  Allocate / Return part
# ============================================================
class PartQtyRequest(BaseModel):
    lot_id: int
    qty: float


def get_or_create_shipment(db: Session, lot: ProductionLot):
    """หรือล็อตนี้มี shipment แล้วหรือยัง — ถ้าไม่มีให้สร้างใหม่"""
    shipment = (
        db.query(CustomerShipment)
        .filter(CustomerShipment.lot_id == lot.id)
        .order_by(CustomerShipment.id.desc())
        .first()
    )

    if not shipment:
        shipment = CustomerShipment(
            po_id=lot.po_id,
            lot_id=lot.id,  # ✅ ผูกตรงกับ lot ปลายทาง
            shipped_at=datetime.now(),
            status="pending",
        )
        db.add(shipment)
        db.flush()

    return shipment



@router.post("/allocate-part")
def allocate_part(req: AllocatePartRequest, db: Session = Depends(get_db)):
    from models import ShopTraveler, ShopTravelerStep

    source_lot = db.get(ProductionLot, req.source_lot_id)
    target_lot = db.get(ProductionLot, req.target_lot_id)

    if not source_lot or not target_lot:
        raise HTTPException(status_code=404, detail="Lot not found")

    qty = float(req.qty)
    if qty <= 0:
        raise HTTPException(status_code=400, detail="Quantity must be greater than zero")

    # --- คำนวณ available จาก source lot ---
    sub_max_seq = (
        db.query(
            ShopTravelerStep.traveler_id,
            func.max(ShopTravelerStep.seq).label("max_seq"),
        )
        .group_by(ShopTravelerStep.traveler_id)
        .subquery()
    )

    finished_qty = (
        db.query(func.coalesce(func.sum(ShopTravelerStep.qty_accept), 0))
        .join(
            sub_max_seq,
            (ShopTravelerStep.traveler_id == sub_max_seq.c.traveler_id)
            & (ShopTravelerStep.seq == sub_max_seq.c.max_seq),
        )
        .join(ShopTraveler, ShopTraveler.id == ShopTravelerStep.traveler_id)
        .filter(ShopTraveler.lot_id == source_lot.id)   # ✅ ใช้ source_lot
        .filter(ShopTravelerStep.status.in_(["passed", "completed"]))
        .scalar()
        or 0
    )
   
    shipped_qty = (
        db.query(func.coalesce(func.sum(CustomerShipmentItem.qty), 0))
        .filter(CustomerShipmentItem.lot_id == source_lot.id)   # ✅ แก้แล้ว
        .scalar()
        or 0
    )

    available_qty = float(finished_qty) - float(shipped_qty)
    if qty > available_qty:
        raise HTTPException(
            status_code=400,
            detail=f"Cannot allocate {qty} pcs — only {available_qty} available in {source_lot.lot_no}",
        )
   
    # ✅ ถ้ามี shipment_id ให้ใช้เลย ไม่ต้องสร้างใหม่
    if req.shipment_id:
        shipment = db.get(CustomerShipment, req.shipment_id)
        if not shipment:
            raise HTTPException(status_code=404, detail="Shipment not found")
    else:
        shipment = get_or_create_shipment(db, target_lot)

    new_item = CustomerShipmentItem(
        shipment_id=shipment.id,
        po_line_id=target_lot.po_line_id or 0,
        lot_id=target_lot.id,  # ของมาจาก lot ต้นทาง
        lot_allocate_id = source_lot.id,
        qty=qty,
    )

    db.add(new_item)
    db.commit()
    db.refresh(new_item)

    return {
        "status": "ok",
        "from": source_lot.lot_no,
        "to": target_lot.lot_no,
        "shipment_id": shipment.id,
        "allocated_qty": qty,
        "available_after": available_qty - qty,
    }


@router.post("/return-part")
def return_part(req: dict = Body(...), db: Session = Depends(get_db)):
    source_lot_id = req.get("source_lot_id")
    target_lot_id = req.get("target_lot_id")
    shipment_id = req.get("shipment_id")  # optional
    qty = float(req.get("qty", 0))

    if not source_lot_id or not target_lot_id:
        raise HTTPException(status_code=400, detail="Missing source_lot_id or target_lot_id")
    if qty <= 0:
        raise HTTPException(status_code=400, detail="Quantity must be > 0")

    # ✅ ดึง shipment ที่ต้องใช้ในการ return (priority: shipment_id)
    if shipment_id:
        shipment = db.get(CustomerShipment, shipment_id)
        if not shipment:
            raise HTTPException(status_code=404, detail="Shipment not found")
        shipments = [shipment]
    else:
        shipments = (
            db.query(CustomerShipment)
            .filter(CustomerShipment.lot_id == target_lot_id)
            .order_by(CustomerShipment.id.desc())
            .all()
        )

    if not shipments:
        raise HTTPException(status_code=400, detail="No shipment found for target lot")

    remain = qty
    total_returned = 0.0

    # ✅ return จาก item ที่ถูก allocate จริงจาก source_lot_id
    for s in shipments:
        items = (
            db.query(CustomerShipmentItem)
            .filter(CustomerShipmentItem.shipment_id == s.id)
            .filter(CustomerShipmentItem.lot_allocate_id == source_lot_id)  # 👈 ต้อง filter จาก allocate_id จริง
            .order_by(CustomerShipmentItem.id.desc())
            .all()
        )

        for it in items:
            if remain <= 0:
                break

            item_qty = float(it.qty or 0)
            if item_qty <= remain:
                remain -= item_qty
                total_returned += item_qty
                total_returned += item_qty
                db.delete(it)
            else:
                it.qty = item_qty - remain
                total_returned += remain
                remain = 0

        if remain <= 0:
            break

    db.commit()

    if total_returned == 0:
        raise HTTPException(status_code=400, detail="No allocatable part qty found to return")

    return {
        "status": "returned",
        "returned_qty": total_returned,
        "source_lot_id": source_lot_id,
        "target_lot_id": target_lot_id,
        "shipment_id": shipment_id or shipments[0].id,
    }



# ============================================================
# 3️⃣  Delete shipment
# ============================================================
@router.delete("/delete/{shipment_id}")
def delete_shipment(shipment_id: int, db: Session = Depends(get_db)):
    s = db.get(CustomerShipment, shipment_id)
    if not s:
        raise HTTPException(status_code=404, detail="Shipment not found")
    db.delete(s)
    db.commit()
    return {"status": "deleted", "shipment_id": shipment_id}

# ============================================================
# 4️⃣  Shipment history
# ============================================================
@router.get("/history/{lot_id}")
def shipment_history(lot_id: int, db: Session = Depends(get_db)):
    rows = (
        db.query(CustomerShipment)
        .join(CustomerShipment.items)
        .filter(CustomerShipmentItem.lot_id == lot_id)
        .order_by(CustomerShipment.shipped_at.desc())
        .all()
    )
    return [
        {
            "id": r.id,
            "shipment_no": r.package_no or f"SHP-{r.id:05d}",
            "qty": sum(float(i.qty or 0) for i in r.items if i.lot_id == lot_id),
            "uom": "pcs",
            "status": r.status or "pending",
            "created_at": r.shipped_at,
        }
        for r in rows
    ]


# ============================================================
# 5️⃣  Header info
# ============================================================
@router.get("/lot/{lot_id}/header")
def get_lot_header(lot_id: int, db: Session = Depends(get_db)):
    lot = (
        db.query(ProductionLot)
        .options(
            joinedload(ProductionLot.part),
            joinedload(ProductionLot.po).joinedload(PO.customer),   # 👈 load PO
        )
        .filter(ProductionLot.id == lot_id)
        .first()
    )
    if not lot:
        raise HTTPException(status_code=404, detail="Lot not found")

    data = get_part_inventory_data(db, lot.id)
    po = lot.po
    customer = po.customer if po else None
    return {
        "lot_id": lot.id,
        "lot_no": lot.lot_no,

        "part_no": data["part"].part_no if data["part"] else None,
        "po_number": data["po_number"],

        # ✅ NEW
        "customer_code": customer.code if customer else None,
        "customer_name": customer.name if customer else None,
        "customer_address": customer.address if customer else None,

        "planned_qty": data["planned_qty"],
        "finished_qty": data["finished_qty"],
        "shipped_qty": data["shipped_qty"],
        "available_qty": data["available_qty"],

        "status": lot.status,
        "due_date": lot.lot_due_date,
    }



# ============================================================
# 6️⃣  Part inventory + progress
# ============================================================

@router.get("/lot/{lot_id}/part-inventory")
def get_part_inventory(lot_id: int, db: Session = Depends(get_db)):
    data = get_part_inventory_data(db, lot_id)

    part = data["part"]
    planned = data["planned_qty"]
    finished = data["finished_qty"]
    shipped = data["shipped_qty"]
    available = data["available_qty"]

    progress_percent = round(finished / planned * 100, 2) if planned > 0 else 0

    return {
        "part_id": part.id,
        "part_no": part.part_no,
        "lot_id": lot_id,
        "planned_qty": planned,
        "finished_qty": finished,
        "shipped_qty": shipped,
        "available_qty": available,
        "progress_percent": progress_percent,
        "uom": getattr(part, "uom", "pcs"),
    }


# ============================================================
# 7️⃣  Update status
# ============================================================
@router.patch("/{shipment_id}/status")
def update_shipment_status(
    shipment_id: int, data: dict = Body(...), db: Session = Depends(get_db)
):
    s = db.get(CustomerShipment, shipment_id)
    if not s:
        raise HTTPException(status_code=404, detail="Shipment not found")

    new_status = data.get("status")
    if new_status not in ["pending", "shipped", "cancelled"]:
        raise HTTPException(status_code=400, detail="Invalid status")

    s.status = new_status
    db.commit()
    return {"status": "ok", "shipment_id": s.id, "new_status": s.status}




# ============================================================
# 6️⃣  Part inventory (ทุก lot ของ part เดียวกัน)
# ============================================================
@router.get("/lot/{lot_id}/part-inventory/all")
def get_part_inventory_all_for_same_part(lot_id: int, db: Session = Depends(get_db)):
    lot = get_lot_or_404(db, lot_id)
    part_id = lot.part_id

    lots = db.query(ProductionLot).filter(ProductionLot.part_id == part_id).all()
    results = []

    for l in lots:
        data = get_part_inventory_data(db, l.id)

        part = data["part"]
        planned = data["planned_qty"]
        finished = data["finished_qty"]
        shipped = data["shipped_qty"]
        available = data["available_qty"]

        progress_percent = round(finished / planned * 100, 2) if planned > 0 else 0

        results.append({
            "lot_id": l.id,
            "lot_no": l.lot_no,
            "part_no": part.part_no,
            "planned_qty": planned,
            "finished_qty": finished,
            "shipped_qty": shipped,
            "available_qty": available,
            "progress_percent": progress_percent,
            "uom": getattr(part, "uom", "pcs"),
        })

    return results

# ============================================================
# 🆕 Create a new shipment for a lot
# ============================================================
@router.post("")
def create_shipment(req: dict = Body(...), db: Session = Depends(get_db)):
    lot_id = req.get("lot_id")
    if not lot_id:
        raise HTTPException(status_code=400, detail="Missing lot_id")

    lot = db.get(ProductionLot, lot_id)
    if not lot:
        raise HTTPException(status_code=404, detail="Lot not found")

    shipment = CustomerShipment(
        po_id=lot.po_id,
        lot_id=lot.id,
        shipped_at=datetime.now(),
        status="pending",
    )
    db.add(shipment)
    db.commit()
    db.refresh(shipment)

    return {
        "status": "created",
        "shipment_id": shipment.id,
        "shipment_no": shipment.package_no or f"SHP-{shipment.id}",
        "lot_no": lot.lot_no,
        "message": f"Shipment created for lot {lot.lot_no}",
    }


@router.patch("/{shipment_id}/update-fields")
def update_shipment_fields(
    shipment_id: int,
    payload: dict = Body(...),
    db: Session = Depends(get_db),
):
    shipment = db.get(CustomerShipment, shipment_id)
    if not shipment:
        raise HTTPException(status_code=404, detail="Shipment not found")

    updated_fields = []

    # ---- Tracking No ----
    # ---- Tracking No ----
    if "tracking_number" in payload:
        shipment.tracking_no = payload["tracking_number"]
        updated_fields.append("tracking_number")

        # Auto-mark as shipped if tracking number is provided
        if payload["tracking_number"]:
            shipment.status = "shipped"
            updated_fields.append("status")

            lot = db.get(ProductionLot, shipment.lot_id)
            if lot:
                lot.status = "completed"
                updated_fields.append("lot.status=completed")

            shipment.shipped_at = datetime.utcnow()
            updated_fields.append("shipped_date")

           

    # ---- Shipped Date ----
    if "shipped_date" in payload:
        try:
            shipment.shipped_at = datetime.fromisoformat(payload["shipped_date"])
        except Exception:
            shipment.shipped_at = datetime.utcnow()
        updated_fields.append("shipped_date")

    if "status" in payload:
        new_status = payload["status"]
        shipment.status = new_status
        updated_fields.append("status")

        # 🔥 If shipment is shipped → mark lot as completed
        if new_status == "shipped":
            lot = db.get(ProductionLot, shipment.lot_id)
            if lot:
                lot.status = "completed"
                updated_fields.append("lot.status=completed")
        if new_status == "pending":
            lot = db.get(ProductionLot, shipment.lot_id)
            if lot:
                lot.status = "not_start"
                updated_fields.append("lot.status=completed")

    # ---- Qty (NEW) ----
    if "qty" in payload:
        # ต้องแก้ที่ CustomerShipmentItem
        items = (
            db.query(CustomerShipmentItem)
            .filter(CustomerShipmentItem.shipment_id == shipment_id)
            .all()
        )
        if items:
            items[0].qty = float(payload["qty"])
            updated_fields.append("qty")

    db.commit()
    db.refresh(shipment)

    return {
        "id": shipment.id,
        "updated_fields": updated_fields,
        "status": shipment.status,
        "tracking_no": shipment.tracking_no,
        "shipped_at": shipment.shipped_at,
    }



# ============================================================
# 8️⃣  Download CofC as DOCX
# ============================================================
from fastapi.responses import FileResponse
from docx import Document
import tempfile
import os

# def generate_docx_from_template(
#     db: Session,
#     shipment_id: int,
#     template_path: str,
#     filename_prefix: str,
#     replace_map_builder,
# ):
#     import os, tempfile
#     from docx import Document
#     from fastapi.responses import FileResponse
#     from datetime import datetime

#     shipment = (
#         db.query(CustomerShipment)
#         .options(
#             joinedload(CustomerShipment.items)
#                 .joinedload(CustomerShipmentItem.lot)
#                 .joinedload(ProductionLot.part),

#             joinedload(CustomerShipment.items)
#                 .joinedload(CustomerShipmentItem.lot)
#                 .joinedload(ProductionLot.part_revision),

#             joinedload(CustomerShipment.po).joinedload(PO.customer),
#         )
#         .get(shipment_id)
#     )

#     if not shipment:
#         raise HTTPException(status_code=404, detail="Shipment not found")

#     if not shipment.items:
#         raise HTTPException(status_code=400, detail="Shipment has no shipment items")

#     item = shipment.items[0]
#     lot = item.lot
#     part = lot.part if lot else None
#     revision = lot.part_revision if lot else None

#     lot_no = lot.lot_no if lot else ""
#     qty = int(item.qty or 0)
#     part_no = part.part_no if part else ""
#     rev = revision.rev if revision else ""
#     desc = part.name if part else ""

#     customer_name = shipment.po.customer.name if shipment.po and shipment.po.customer else ""
#     customer_address = shipment.po.customer.address if shipment.po and shipment.po.customer else ""
#     po_no = shipment.po.po_number if shipment.po else ""

#     fair = lot.fair_note if part else ""

#     # Let caller decide what fields to use
#     replace_map = replace_map_builder(
#         lot_no, part_no, rev, qty, desc,
#         customer_name, customer_address, po_no, shipment.id, fair
#     )

#     if not os.path.exists(template_path):
#         raise HTTPException(status_code=404, detail="Template not found")

#     doc = Document(template_path)

#     def replace_runs(paragraph):
#         for run in paragraph.runs:
#             for k, v in replace_map.items():
#                 if k in run.text:
#                     run.text = run.text.replace(k, v or "")

#     for p in doc.paragraphs:
#         replace_runs(p)

#     for table in doc.tables:
#         for row in table.rows:
#             for cell in row.cells:
#                 for p in cell.paragraphs:
#                     replace_runs(p)

#     download_name = f"{filename_prefix}_{lot_no}_{part_no}.docx"

#     tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
#     doc.save(tmp.name)

#     return FileResponse(
#         tmp.name,
#         media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
#         filename=download_name,
#     )

def generate_docx_from_template(
    db: Session,
    shipment_id: int,
    template_path: str,
    filename_prefix: str,
    replace_map_builder,
):
    import os, tempfile
    from docx import Document
    from fastapi.responses import FileResponse
    from datetime import datetime

    shipment = (
        db.query(CustomerShipment)
        .options(
            # LOT info
            joinedload(CustomerShipment.items)
                .joinedload(CustomerShipmentItem.lot)
                .joinedload(ProductionLot.part),

            joinedload(CustomerShipment.items)
                .joinedload(CustomerShipmentItem.lot)
                .joinedload(ProductionLot.part_revision),

            # ✅ NEW: PO LINE REV
            joinedload(CustomerShipment.items)
                .joinedload(CustomerShipmentItem.po_line)
                .joinedload(POLine.rev),

            joinedload(CustomerShipment.po).joinedload(PO.customer),
        )
        .get(shipment_id)
    )

    if not shipment:
        raise HTTPException(status_code=404, detail="Shipment not found")

    if not shipment.items:
        raise HTTPException(status_code=400, detail="Shipment has no shipment items")

    item = shipment.items[0]

    lot = item.lot
    part = lot.part if lot else None

    # =========================
    # ✅ FIXED REV FROM PO LINE
    # =========================
    rev = ""
    for i in shipment.items:
        if i.po_line and i.po_line.rev:
            rev = i.po_line.rev.rev
            break

    # =========================
    # NORMAL FIELDS
    # =========================
    lot_no = lot.lot_no if lot else ""
    qty = int(item.qty or 0)
    part_no = part.part_no if part else ""
    desc = part.name if part else ""

    customer_name = shipment.po.customer.name if shipment.po and shipment.po.customer else ""
    customer_address = shipment.po.customer.address if shipment.po and shipment.po.customer else ""
    po_no = shipment.po.po_number if shipment.po else ""

    fair = lot.fair_note if lot else ""

    replace_map = replace_map_builder(
        lot_no, part_no, rev, qty, desc,
        customer_name, customer_address, po_no, shipment.id, fair
    )

    if not os.path.exists(template_path):
        raise HTTPException(status_code=404, detail="Template not found")

    doc = Document(template_path)

    def replace_runs(paragraph):
        for run in paragraph.runs:
            for k, v in replace_map.items():
                if k in run.text:
                    run.text = run.text.replace(k, v or "")

    for p in doc.paragraphs:
        replace_runs(p)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_runs(p)

    download_name = f"{filename_prefix}_{lot_no}_{part_no}.docx"

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name)

    return FileResponse(
        tmp.name,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=download_name,
    )

@router.get("/{shipment_id}/download/cofc")
def download_cofc(shipment_id: int, db: Session = Depends(get_db)):
    print(f"Generating CofC for shipment {shipment_id}")

    def build_map(lot_no, part_no, rev, qty, desc, cust, addr, po, sid,fair):
        from datetime import datetime
        return {
            "{CUSTOMER}": cust,
            "{CUSTOMER_ADDRESS}": addr,
            "{PO_NO}": po,
            "{PART_NO}": part_no,
            "{REV}": rev,
            "{QTY}": str(qty),
            "{LOT_NO}": lot_no,
            "{LOT}": lot_no,
            "{DESCRIPTION}": desc,
            "{CERT_NO}": f"CERT-{sid:05d}",
            "{DATE}": datetime.now().strftime("%m/%d/%Y"),
            
        }

    return generate_docx_from_template(
        db=db,
        shipment_id=shipment_id,
        template_path="templates/cofc.docx",
        filename_prefix="cofc",
        replace_map_builder=build_map,
    )

@router.get("/{shipment_id}/download/packing")
def download_packing(shipment_id: int, db: Session = Depends(get_db)):

    def build_map(lot_no, part_no, rev, qty, desc, cust, addr, po, sid, fair):
        return {
            "{CUSTOMER}": cust,
            "{CUSTOMER_ADDRESS}": addr,
            "{PO_NO}": po,
            "{PART_NO}": part_no,
            "{REV}": rev,
            "{QTY}": str(qty),
            "{LOT_NO}": lot_no,
            "{LOT}": lot_no,
            "{DESCRIPTION}": desc,
        }

    return generate_docx_from_template(
        db=db,
        shipment_id=shipment_id,
        template_path="templates/packing.docx",
        filename_prefix="packing",
        replace_map_builder=build_map,
    )


@router.get("/{shipment_id}/download/packingFA")
def download_packing(shipment_id: int, db: Session = Depends(get_db)):

    def build_map(lot_no, part_no, rev, qty, desc, cust, addr, po, sid,fair):
        return {
            "{CUSTOMER}": cust,
            "{CUSTOMER_ADDRESS}": addr,
            "{PO_NO}": po,
            "{PART_NO}": part_no,
            "{REV}": rev,
            "{QTY}": str(qty),
            "{LOT_NO}": lot_no,
            "{LOT}": lot_no,
            "{DESCRIPTION}": desc,
            "{FAIR}" : fair
        }

    return generate_docx_from_template(
        db=db,
        shipment_id=shipment_id,
        template_path="templates/packing.docx",
        filename_prefix="packing",
        replace_map_builder=build_map,
    )


##Labels: recent-edits
from fastapi import Query
from fastapi.responses import FileResponse
import tempfile
from docx import Document
from datetime import datetime
import os

from fastapi import Query, Depends, HTTPException
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session, joinedload
from docx import Document
from datetime import datetime
import tempfile
import os

from docx.oxml.ns import qn


# @router.get("/{shipment_id}/download/label/{size}")
# def download_label(
#     shipment_id: int,
#     size: int,
#     type: str = Query(None, description="Type of label"),
#     db: Session = Depends(get_db),
# ):
#     print(f"Generating label for shipment {shipment_id}, size {size}, type {type}")

#     # ===============================
#     # 1️⃣ LOAD DATA
#     # ===============================
#     shipment = (
#         db.query(CustomerShipment)
#         .options(
#             joinedload(CustomerShipment.po).joinedload(PO.customer),
#             joinedload(CustomerShipment.items)
#                 .joinedload(CustomerShipmentItem.lot)
#                 .joinedload(ProductionLot.part),
#             joinedload(CustomerShipment.items)
#                 .joinedload(CustomerShipmentItem.lot)
#                 .joinedload(ProductionLot.part_revision),
#         )
#         .filter(CustomerShipment.id == shipment_id)
#         .first()
#     )

#     if not shipment or not shipment.items:
#         raise HTTPException(status_code=404, detail="Shipment not found")

#     item = shipment.items[0]
#     lot = item.lot
#     part = lot.part if lot else None
#     revision = lot.part_revision if lot else None

#     lot_no = lot.lot_no if lot else "UNKNOWN_LOT"
#     part_no = part.part_no if part else "UNKNOWN_PART"
#     po_no = shipment.po.po_number if shipment.po else ""

#     customer_name = (
#         shipment.po.customer.name
#         if shipment.po and shipment.po.customer
#         else ""
#     )

#     # ===============================
#     # 2️⃣ QTY LOGIC
#     # ===============================
   

#     if type == "fair":
#         total_qty = 1
#     elif type == "cmm":
#          total_qty = int(item.qty or 0) 
#     elif type == "number":
#         total_qty = int(item.qty or 0) 
#     elif type == "box":
#         total_qty = int(item.qty or 0) 
#     else:
#          total_qty = int(item.qty or 0) + 2

#     page_size = size

#     if total_qty <= page_size:
#         first_page_limit = total_qty
#         second_page_limit = 0
#         remove_page2 = True
#     else:
#         first_page_limit = page_size
#         second_page_limit = total_qty % page_size
#         remove_page2 = False
# # 
#     # print(f"[DEBUG] total_qty={total_qty}")
#     # print(f"[DEBUG] page1={first_page_limit}, page2={second_page_limit}")

#     # ===============================
#     # 3️⃣ REPLACE MAP
#     # ===============================
#     replace_map = {
#         "{PART}": f"Part: {part.part_no} {part.name}" if part else "",
#         "{REV}": f"Rev: {revision.rev} Lot: {lot.lot_no} PO: {po_no}" if revision else "",
#         "{LOT_NO}": lot_no,
#         "{DESCRIPTION}": part.name if part else "",
#         "{DATE}": datetime.now().strftime("%m/%d/%Y"),
#         "CUSTOMER": customer_name,
#         "{XX}": total_qty,
#     }

#     # ===============================
#     # 4️⃣ LOAD TEMPLATE
#     # ===============================
#     print(total_qty , size)
#     if type == "fair":
#         template_path = "templates/label_fair.docx"
#     elif type == "cmm":
#         template_path = "templates/label_cmm.docx"
#     elif type == "number":
#         template_path = "templates/label_number.docx"
#     elif type == "box":
#         template_path = "templates/label_box.docx"
#     else:
#         if total_qty < size:
#             print(f"[DEBUG] Using 1-page template for size {size} because total_qty={total_qty} < page_size={page_size}")
#             template_path = f"templates/label_{size}_1_page.docx"
#         else:
#             template_path = f"templates/label_{size}.docx"

#     if not os.path.exists(template_path):
#         raise HTTPException(status_code=404, detail="Template not found")

#     doc = Document(template_path)

#     # ===============================
#     # 5️⃣ REPLACE FUNCTIONS
#     # ===============================
#     def replace_runs(p):
#         full_text = "".join(run.text for run in p.runs)
#         original = full_text

#         for k, v in replace_map.items():
#             if k in full_text:
#                 full_text = full_text.replace(k, str(v))

#         if full_text != original:
#             for run in p.runs:
#                 run.text = ""
#             if p.runs:
#                 p.runs[0].text = full_text

                

#     def clear_runs(p):
#         for run in p.runs:
#             run.text = ""

#     # ===============================
#     # 6️⃣ GLOBAL LABEL CONTROL
#     # ===============================
#     label_index = 0

#     def process_paragraph(p):
#         nonlocal label_index

#         text = p.text.strip()

#         if "Topnotch Quality Works" in text:
#             label_index += 1
#             # print(f"[DEBUG] Label #{label_index}")

#         # CASE 1 → ONLY ONE PAGE
#         if remove_page2:
#             if label_index <= first_page_limit:
#                 replace_runs(p)
#             else:
#                 clear_runs(p)
#             return

#         # CASE 2 → TWO PAGE LOGIC
#         if label_index <= first_page_limit:
#             replace_runs(p)

#         elif label_index <= first_page_limit + second_page_limit:
#             replace_runs(p)

#         else:
#             clear_runs(p)

#     # ===============================
#     # 7️⃣ PROCESS CONTENT
#     # ===============================
#     for p in doc.paragraphs:
#         process_paragraph(p)

#     for table in doc.tables:
#         for row in table.rows:
#             for cell in row.cells:
#                 for p in cell.paragraphs:
#                     process_paragraph(p)

#     # ===============================
#     # 8️⃣ REMOVE PAGE BREAK (IMPORTANT)
#     # ===============================
#     if remove_page2:
#         for paragraph in doc.paragraphs:
#             p = paragraph._element
#             for br in p.findall(".//w:br", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}):
#                 if br.get(qn("w:type")) == "page":
#                     p.remove(br)

#     # ===============================
#     # 9️⃣ SAVE FILE
#     # ===============================
#     tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
#     doc.save(tmp.name)

#     # ===============================
#     # 🔟 RETURN FILE
#     # ===============================
#     label_type = type.upper() if type else "LABEL"
#     filename = f"label_{lot_no}_{part_no}_{label_type}_{size}.docx"

#     return FileResponse(
#         path=tmp.name,
#         media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
#         filename=filename,
#     )

@router.get("/{shipment_id}/download/label/{size}")
def download_label(
    shipment_id: int,
    size: int,
    type: str = Query(None, description="Type of label"),
    db: Session = Depends(get_db),
):
    print(f"Generating label for shipment {shipment_id}, size {size}, type {type}")

    # ===============================
    # 1️⃣ LOAD DATA
    # ===============================
    shipment = (
        db.query(CustomerShipment)
        .options(
            joinedload(CustomerShipment.po).joinedload(PO.customer),
            joinedload(CustomerShipment.items)
                .joinedload(CustomerShipmentItem.lot)
                .joinedload(ProductionLot.part),
            joinedload(CustomerShipment.items)
                .joinedload(CustomerShipmentItem.lot)
                .joinedload(ProductionLot.part_revision),
            joinedload(CustomerShipment.items)
                .joinedload(CustomerShipmentItem.po_line)
                .joinedload(POLine.rev),
        )
        .filter(CustomerShipment.id == shipment_id)
        .first()
    )

    if not shipment or not shipment.items:
        raise HTTPException(status_code=404, detail="Shipment not found")

    item = shipment.items[0]
    lot = item.lot
    part = lot.part if lot else None
    rev = ""
    for i in shipment.items:
        if i.po_line and i.po_line.rev:
            rev = i.po_line.rev.rev
            break

    lot_no = lot.lot_no if lot else "UNKNOWN_LOT"
    part_no = part.part_no if part else "UNKNOWN_PART"
    po_no = shipment.po.po_number if shipment.po else ""

    customer_name = (
        shipment.po.customer.name
        if shipment.po and shipment.po.customer
        else ""
    )

    # ===============================
    # 2️⃣ QTY LOGIC
    # ===============================
    if type == "fair":
        total_qty = 1
    elif type in ["cmm", "number", "box"]:
        total_qty = int(item.qty or 0)
    else:
        total_qty = int(item.qty or 0) + 2

    page_size = size

    if total_qty <= page_size:
        first_page_limit = total_qty
        second_page_limit = 0
        remove_page2 = True
    else:
        first_page_limit = page_size
        second_page_limit = total_qty % page_size
        remove_page2 = False

    # ===============================
    # 3️⃣ REPLACE MAP
    # ===============================
    replace_map = {
        "{PART}": f"Part: {part.part_no} {part.name}" if part else "",
        "{REV}": f"Rev: {rev} Lot: {lot.lot_no} PO: {po_no}" if rev else "",
        "{LOT_NO}": lot_no,
        "{DESCRIPTION}": part.name if part else "",
        "{DATE}": datetime.now().strftime("%m/%d/%Y"),
        "CUSTOMER": customer_name,
    }

    # ===============================
    # 4️⃣ LOAD TEMPLATE
    # ===============================
    if type == "fair":
        template_path = "templates/label_fair.docx"
    elif type == "cmm":
        template_path = "templates/label_cmm.docx"
    elif type == "number":
        template_path = "templates/label_number.docx"
    elif type == "box":
        template_path = "templates/label_box.docx"
    else:
        if total_qty < size:
            template_path = f"templates/label_{size}_1_page.docx"
        else:
            template_path = f"templates/label_{size}.docx"

    if not os.path.exists(template_path):
        raise HTTPException(status_code=404, detail="Template not found")

    doc = Document(template_path)
    print(f"[DEBUG] Loaded template: {template_path}")

    # ===============================
    # 5️⃣ REPLACE FUNCTIONS
    # ===============================
    def replace_runs(p, current_index=None):
        full_text = "".join(run.text for run in p.runs)
        original = full_text

        for k, v in replace_map.items():
            if k in full_text:
                full_text = full_text.replace(k, str(v))

        if "{XX}" in full_text and current_index is not None:
            full_text = full_text.replace("{XX}", str(current_index))

        if full_text != original:
            for run in p.runs:
                run.text = ""
            if p.runs:
                p.runs[0].text = full_text

    def clear_runs(p):
        for run in p.runs:
            run.text = ""

    # ===============================
    # 6️⃣ NORMAL PARAGRAPHS
    # ===============================
    label_index = 0

    for p in doc.paragraphs:
        text = p.text.strip()

        if "Topnotch Quality Works" in text:
            label_index += 1

        if remove_page2:
            if label_index <= first_page_limit:
                replace_runs(p, label_index)
            else:
                clear_runs(p)
            continue

        if label_index <= first_page_limit:
            replace_runs(p, label_index)

        elif label_index <= first_page_limit + second_page_limit:
            replace_runs(p, label_index)

        else:
            clear_runs(p)

    # ===============================
    # 7️⃣ TABLE PROCESS (SMART ORDER)
    # ===============================

    for table in doc.tables:
        rows = len(table.rows)
        cols = len(table.rows[0].cells)
        print(f"[DEBUG] Processing table with {rows} rows and {cols} cols")
        ordered_cells = []

        if total_qty > size:
            # ✅ row-first
           for r in range(rows):
                for c in range(len(table.rows[r].cells)):
                    ordered_cells.append(table.rows[r].cells[c])
        else:
            # ✅ column-first
             for r in range(rows):
                for c in range(len(table.rows[r].cells)):
                    ordered_cells.append(table.rows[r].cells[c])

        label_index = 0

        for cell in ordered_cells:
            for p in cell.paragraphs:
                text = p.text.strip()

                if "Topnotch Quality Works" in text:
                    label_index += 1

                if remove_page2 :
                    if label_index <= first_page_limit:
                        replace_runs(p, label_index)
                    else:
                        clear_runs(p)
                    continue

                if label_index <= first_page_limit:
                    replace_runs(p, label_index)

                elif label_index <= first_page_limit + second_page_limit:
                    replace_runs(p, label_index)

                else:
                    clear_runs(p)

    # ===============================
    # 8️⃣ REMOVE PAGE BREAK
    # ===============================
    if remove_page2:
        for paragraph in doc.paragraphs:
            p = paragraph._element
            for br in p.findall(
                ".//w:br",
                namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"},
            ):
                if br.get(qn("w:type")) == "page":
                    p.remove(br)

    # ===============================
    # 9️⃣ SAVE FILE
    # ===============================
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name)

    # ===============================
    # 🔟 RETURN FILE
    # ===============================
    label_type = type.upper() if type else "LABEL"
    filename = f"label_{lot_no}_{part_no}_{label_type}_{size}.docx"

    return FileResponse(
        path=tmp.name,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename,
    )

# @router.get("/{shipment_id}/download/label/{size}")
# def download_label(
#     shipment_id: int,
#     size: int,
#     type: str = Query(None, description="Type of label"),
#     db: Session = Depends(get_db),
# ):
#     print(f"Generating label for shipment {shipment_id}, size {size}, type {type}")

#     # ===============================
#     # 1️⃣ LOAD DATA
#     # ===============================
#     shipment = (
#         db.query(CustomerShipment)
#         .options(
#             joinedload(CustomerShipment.po).joinedload(PO.customer),
#             joinedload(CustomerShipment.items)
#                 .joinedload(CustomerShipmentItem.lot)
#                 .joinedload(ProductionLot.part),
#             joinedload(CustomerShipment.items)
#                 .joinedload(CustomerShipmentItem.lot)
#                 .joinedload(ProductionLot.part_revision),
#         )
#         .filter(CustomerShipment.id == shipment_id)
#         .first()
#     )

#     if not shipment or not shipment.items:
#         raise HTTPException(status_code=404, detail="Shipment not found")

#     item = shipment.items[0]
#     lot = item.lot
#     part = lot.part if lot else None
#     revision = lot.part_revision if lot else None

#     lot_no = lot.lot_no if lot else "UNKNOWN_LOT"
#     part_no = part.part_no if part else "UNKNOWN_PART"
#     po_no = shipment.po.po_number if shipment.po else ""

#     customer_name = (
#         shipment.po.customer.name
#         if shipment.po and shipment.po.customer
#         else ""
#     )

#     # ===============================
#     # 2️⃣ QTY LOGIC
#     # ===============================
#     if type == "fair":
#         total_qty = 1
#     elif type in ["cmm", "number", "box"]:
#         total_qty = int(item.qty or 0)
#     else:
#         total_qty = int(item.qty or 0) + 2

#     page_size = size

#     if total_qty <= page_size:
#         first_page_limit = total_qty
#         second_page_limit = 0
#         remove_page2 = True
#     else:
#         first_page_limit = page_size
#         second_page_limit = total_qty % page_size
#         remove_page2 = False

#     # ===============================
#     # 3️⃣ REPLACE MAP
#     # ===============================
#     replace_map = {
#         "{PART}": f"Part: {part.part_no} {part.name}" if part else "",
#         "{REV}": f"Rev: {revision.rev} Lot: {lot.lot_no} PO: {po_no}" if revision else "",
#         "{LOT_NO}": lot_no,
#         "{DESCRIPTION}": part.name if part else "",
#         "{DATE}": datetime.now().strftime("%m/%d/%Y"),
#         "CUSTOMER": customer_name,
#     }

#     # ===============================
#     # 4️⃣ LOAD TEMPLATE
#     # ===============================
#     if type == "fair":
#         template_path = "templates/label_fair.docx"
#     elif type == "cmm":
#         template_path = "templates/label_cmm.docx"
#     elif type == "number":
#         template_path = "templates/label_number.docx"
#     elif type == "box":
#         template_path = "templates/label_box.docx"
#     else:
#         if total_qty < size:
#             template_path = f"templates/label_{size}_1_page.docx"
#         else:
#             template_path = f"templates/label_{size}.docx"

#     if not os.path.exists(template_path):
#         raise HTTPException(status_code=404, detail="Template not found")

#     doc = Document(template_path)

#     # ===============================
#     # 5️⃣ REPLACE FUNCTIONS
#     # ===============================
#     def replace_runs(p, current_index=None):
#         full_text = "".join(run.text for run in p.runs)
#         original = full_text

#         # normal replace
#         for k, v in replace_map.items():
#             if k in full_text:
#                 full_text = full_text.replace(k, str(v))

#         # 🔥 dynamic numbering
#         if "{XX}" in full_text and current_index is not None:
#             full_text = full_text.replace("{XX}", str(current_index))

#         if full_text != original:
#             for run in p.runs:
#                 run.text = ""
#             if p.runs:
#                 p.runs[0].text = full_text

#     def clear_runs(p):
#         for run in p.runs:
#             run.text = ""

#     # ===============================
#     # 6️⃣ PROCESS NORMAL PARAGRAPHS
#     # ===============================
#     label_index = 0

#     for p in doc.paragraphs:
#         text = p.text.strip()

#         if "Topnotch Quality Works" in text:
#             label_index += 1

#         if remove_page2:
#             if label_index <= first_page_limit:
#                 replace_runs(p, label_index)
#             else:
#                 clear_runs(p)
#             continue

#         if label_index <= first_page_limit:
#             replace_runs(p, label_index)

#         elif label_index <= first_page_limit + second_page_limit:
#             replace_runs(p, label_index)

#         else:
#             clear_runs(p)

#     # ===============================
#     # 7️⃣ PROCESS TABLE (COLUMN-FIRST)
#     # ===============================
#     for table in doc.tables:
#         rows = len(table.rows)
#         cols = len(table.rows[0].cells)

#         # column-first order
#         ordered_cells = []
#         for c in range(cols):
#             for r in range(rows):
#                 ordered_cells.append(table.rows[r].cells[c])

#         label_index = 0  # reset for table

#         for cell in ordered_cells:
#             for p in cell.paragraphs:
#                 text = p.text.strip()

#                 if "Topnotch Quality Works" in text:
#                     label_index += 1

#                 if remove_page2:
#                     if label_index <= first_page_limit:
#                         replace_runs(p, label_index)
#                     else:
#                         clear_runs(p)
#                     continue

#                 if label_index <= first_page_limit:
#                     replace_runs(p, label_index)

#                 elif label_index <= first_page_limit + second_page_limit:
#                     replace_runs(p, label_index)

#                 else:
#                     clear_runs(p)

#     # ===============================
#     # 8️⃣ REMOVE PAGE BREAK
#     # ===============================
#     if remove_page2:
#         for paragraph in doc.paragraphs:
#             p = paragraph._element
#             for br in p.findall(".//w:br", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}):
#                 if br.get(qn("w:type")) == "page":
#                     p.remove(br)

#     # ===============================
#     # 9️⃣ SAVE FILE
#     # ===============================
#     tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
#     doc.save(tmp.name)

#     # ===============================
#     # 🔟 RETURN FILE
#     # ===============================
#     label_type = type.upper() if type else "LABEL"
#     filename = f"label_{lot_no}_{part_no}_{label_type}_{size}.docx"

#     return FileResponse(
#         path=tmp.name,
#         media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
#         filename=filename,
#     )

# from fastapi import Query, Depends, HTTPException
# from fastapi.responses import FileResponse
# from sqlalchemy.orm import Session, joinedload
# from docx import Document
# from datetime import datetime
# import tempfile
# import os
# from docx.oxml.ns import qn


# @router.get("/{shipment_id}/download/label/{size}")
# def download_label(
#     shipment_id: int,
#     size: int,
#     type: str = Query(None, description="Type of label"),
#     db: Session = Depends(get_db),
# ):
#     print(f"Generating label for shipment {shipment_id}, size {size}, type {type}")

#     # ===============================
#     # 1️⃣ LOAD DATA
#     # ===============================
#     shipment = (
#         db.query(CustomerShipment)
#         .options(
#             joinedload(CustomerShipment.po).joinedload(PO.customer),
#             joinedload(CustomerShipment.items)
#                 .joinedload(CustomerShipmentItem.lot)
#                 .joinedload(ProductionLot.part),
#             joinedload(CustomerShipment.items)
#                 .joinedload(CustomerShipmentItem.lot)
#                 .joinedload(ProductionLot.part_revision),
#         )
#         .filter(CustomerShipment.id == shipment_id)
#         .first()
#     )

#     if not shipment or not shipment.items:
#         raise HTTPException(status_code=404, detail="Shipment not found")

#     item = shipment.items[0]
#     lot = item.lot
#     part = lot.part if lot else None
#     revision = lot.part_revision if lot else None

#     lot_no = lot.lot_no if lot else "UNKNOWN_LOT"
#     part_no = part.part_no if part else "UNKNOWN_PART"
#     po_no = shipment.po.po_number if shipment.po else ""

#     customer_name = (
#         shipment.po.customer.name
#         if shipment.po and shipment.po.customer
#         else ""
#     )

#     # ===============================
#     # 2️⃣ QTY LOGIC
#     # ===============================
#     total_qty = int(item.qty or 0) + 2
#     page_size = size

#     if total_qty <= page_size:
#         first_page_limit = total_qty
#         second_page_limit = 0
#         remove_page2 = True
#     else:
#         first_page_limit = page_size
#         second_page_limit = total_qty % page_size
#         remove_page2 = False

#     print(f"[DEBUG] total_qty={total_qty}")
#     print(f"[DEBUG] page1={first_page_limit}, page2={second_page_limit}")

#     # ===============================
#     # 3️⃣ REPLACE MAP
#     # ===============================
#     replace_map = {
#         "{PART}": f"Part: {part.part_no} {part.name}" if part else "",
#         "{REV}": f"Rev: {revision.rev} Lot: {lot.lot_no} PO: {po_no}" if revision else "",
#         "{LOT_NO}": lot_no,
#         "{DESCRIPTION}": part.name if part else "",
#         "{DATE}": datetime.now().strftime("%m/%d/%Y"),
#         "CUSTOMER": customer_name,
#         "{XX}": total_qty,
#     }

#     # ===============================
#     # 4️⃣ LOAD TEMPLATE
#     # ===============================
#     if type == "fair":
#         template_path = "templates/label_fair.docx"
#     elif type == "cmm":
#         template_path = "templates/label_cmm.docx"
#     elif type == "number":
#         template_path = "templates/label_number.docx"
#     elif type == "box":
#         template_path = "templates/label_box.docx"
#     else:
#         template_path = f"templates/label_{size}.docx"

#     if not os.path.exists(template_path):
#         raise HTTPException(status_code=404, detail="Template not found")

#     doc = Document(template_path)

#     # ===============================
#     # 5️⃣ REPLACE FUNCTIONS
#     # ===============================
#     def replace_runs(p):
#         full_text = "".join(run.text for run in p.runs)

#         for k, v in replace_map.items():
#             if k in full_text:
#                 full_text = full_text.replace(k, str(v))

#         for run in p.runs:
#             run.text = ""
#         if p.runs:
#             p.runs[0].text = full_text

#     def clear_runs(p):
#         for run in p.runs:
#             run.text = ""

#     # ===============================
#     # 6️⃣ COLLECT LABEL BLOCKS
#     # ===============================
#     labels = []
#     current_block = []

#     for table in doc.tables:
#         rows = len(table.rows)
#         cols = len(table.rows[0].cells)

#         # 🔥 COLUMN-FIRST (สำคัญ)
#         for c in range(cols):
#             for r in range(rows):
#                 cell = table.cell(r, c)

#                 for p in cell.paragraphs:
#                     text = p.text.strip()

#                     if "Topnotch Quality Works" in text:
#                         if current_block:
#                             labels.append(current_block)
#                         current_block = [p]
#                     else:
#                         if current_block:
#                             current_block.append(p)

#     if current_block:
#         labels.append(current_block)

#     print(f"[DEBUG] total labels found = {len(labels)}")

#     # ===============================
#     # 7️⃣ APPLY PAGE LOGIC (🔥 FINAL FIX)
#     # ===============================
#     for i, block in enumerate(labels, start=1):

#         # ---------- PAGE 1 ----------
#         if i <= first_page_limit:
#             should_fill = True

#         # ---------- PAGE 2 ----------
#         elif i <= first_page_limit + second_page_limit:
#             page2_index = i - first_page_limit  # 🔥 RESET INDEX
#             should_fill = page2_index <= second_page_limit

#         # ---------- EMPTY ----------
#         else:
#             should_fill = False

#         print(f"[DEBUG] label {i} -> fill={should_fill}")

#         for p in block:
#             if should_fill:
#                 replace_runs(p)
#             else:
#                 clear_runs(p)

#     # ===============================
#     # 8️⃣ REMOVE PAGE BREAK
#     # ===============================
#     if remove_page2:
#         for paragraph in doc.paragraphs:
#             p = paragraph._element
#             for br in p.findall(
#                 ".//w:br",
#                 namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"},
#             ):
#                 if br.get(qn("w:type")) == "page":
#                     p.remove(br)

#     # ===============================
#     # 9️⃣ SAVE FILE
#     # ===============================
#     tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
#     doc.save(tmp.name)

#     # ===============================
#     # 🔟 RETURN FILE
#     # ===============================
#     label_type = type.upper() if type else "LABEL"
#     filename = f"label_{lot_no}_{part_no}_{label_type}_{size}.docx"

#     return FileResponse(
#         path=tmp.name,
#         media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
#         filename=filename,
#     )