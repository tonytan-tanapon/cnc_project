from fastapi import APIRouter, Depends, Query
from sqlalchemy.orm import Session
from sqlalchemy import or_

from database import get_db

from models import (
ProductionLot,
LotMaterialUse,
RawBatch,
RawMaterial,
Supplier,
Part,
)

router = APIRouter(
prefix="/material-traceability",
tags=["material-traceability"]
)

@router.get("")
def get_material_traceability(
q: str | None = Query(None),
db: Session = Depends(get_db)
):
    rows = (
        db.query(
            LotMaterialUse.id.label("lot_material_use_id"),
            RawBatch.id.label("batch_id"),
            ProductionLot.id.label("lot_id"),
            ProductionLot.lot_no.label("lot_no"),

            Part.part_no.label("part_no"),

            RawBatch.mat_po.label("mat_po"),
            RawBatch.cutting_note.label("cutting_note"),
            RawBatch.po_note.label("po_note"),

            Supplier.name.label("supplier_name"),

            RawMaterial.type.label("material_type"),
            RawMaterial.spec.label("material_spec"),

            RawBatch.heat_lot.label("heat_lot"),
            RawBatch.size_text.label("size_text"),
            RawBatch.length_text.label("length_text"),

            RawBatch.batch_no.label("batch_no"),
        )

        .join(
            LotMaterialUse,
            LotMaterialUse.lot_id == ProductionLot.id
        )

        .join(
            RawBatch,
            RawBatch.id == LotMaterialUse.batch_id
        )

        .join(
            RawMaterial,
            RawMaterial.id == RawBatch.material_id
        )

        .outerjoin(
            Supplier,
            Supplier.id == RawBatch.supplier_id
        )

        .join(
            Part,
            Part.id == ProductionLot.part_id
        )
    )

    if q:

        search = f"%{q}%"

        rows = rows.filter(
            or_(
                ProductionLot.lot_no.ilike(search),

                Part.part_no.ilike(search),

                RawBatch.batch_no.ilike(search),
                RawBatch.heat_lot.ilike(search),
                RawBatch.mat_po.ilike(search),

                RawBatch.size_text.ilike(search),
                RawBatch.length_text.ilike(search),

                RawBatch.cutting_note.ilike(search),
                RawBatch.po_note.ilike(search),

                Supplier.name.ilike(search),

                RawMaterial.type.ilike(search),
                RawMaterial.spec.ilike(search),
            )
        )

    rows = (
        rows
        .order_by(
            ProductionLot.id.desc()
        )
        .all()
    )

    return [
        dict(r._mapping)
        for r in rows
    ]

from datetime import date
from pydantic import BaseModel
from fastapi import HTTPException

class ProductionLotUpdate(BaseModel):
    lot_no: str | None = None
    planned_qty: float | None = None
    planned_ship_qty: float | None = None
    lot_due_date: date | None = None
    note: str | None = None
    fair_note: str | None = None


@router.put("/production-lots/{lot_id}")
def update_lot(
    lot_id: int,
    payload: ProductionLotUpdate,
    db: Session = Depends(get_db)
):
    
    print("UPDATE LOT")
    print("LOT ID =", lot_id)
    print("PAYLOAD =", payload)
    lot = db.get(ProductionLot, lot_id)

    if not lot:
        raise HTTPException(
            status_code=404,
            detail="Lot not found"
        )

    data = payload.model_dump(exclude_unset=True)

    for key, value in data.items():
        setattr(lot, key, value)

    db.commit()
    db.refresh(lot)

    return {
        "ok": True,
        "lot_id": lot.id
    }

from pydantic import BaseModel

class LotMaterialUseUpdate(BaseModel):
    batch_id: int

@router.put("/lot-material-use/{id}")
def update_lot_material_use(
    id: int,
    payload: LotMaterialUseUpdate,
    db: Session = Depends(get_db)
):
    row = db.get(LotMaterialUse, id)

    if not row:
        raise HTTPException(
            status_code=404,
            detail="LotMaterialUse not found"
        )

    batch = db.get(
        RawBatch,
        payload.batch_id
    )

    row.batch_id = payload.batch_id
    row.raw_material_id = batch.material_id

    db.commit()

    return {"ok": True}


@router.get("/batch-options")
def get_batch_options(
    db: Session = Depends(get_db)
):
    rows = (
        db.query(
            RawBatch.id,
            RawBatch.batch_no,
            RawBatch.heat_lot,

            RawMaterial.type,
            RawMaterial.spec,
        )

        .join(
            RawMaterial,
            RawMaterial.id == RawBatch.material_id
        )

        .order_by(
            RawBatch.batch_no
        )

        .all()
    )

    return [
        {
            "value": r.id,

            "label":
                f"{r.batch_no}"
                f" | {r.type or ''}"
                f" | {r.spec or ''}"
        }
        for r in rows
    ]

class MaterialTraceabilityCreate(
    BaseModel
):
    lot_id: int
    batch_id: int
    qty: float = 0


@router.post("")
def create_material_traceability(
    payload: MaterialTraceabilityCreate,
    db: Session = Depends(get_db)
):

    batch = db.get(
        RawBatch,
        payload.batch_id
    )

    row = LotMaterialUse(
        lot_id=payload.lot_id,
        batch_id=payload.batch_id,
        raw_material_id=batch.material_id,
        qty=payload.qty,
    )

    db.add(row)
    db.commit()
    db.refresh(row)

    return {
        "ok": True
    }

@router.get("/lot-options")
def get_lot_options(
    db: Session = Depends(get_db)
):
    rows = (
        db.query(
            ProductionLot.id,
            ProductionLot.lot_no,
        )
        .order_by(
            ProductionLot.lot_no
        )
        .all()
    )

    return [
        {
            "value": r.id,
            "label": r.lot_no
        }
        for r in rows
    ]

@router.delete("/lot-material-use/{id}")
def delete_lot_material_use(
    id: int,
    db: Session = Depends(get_db)
):
    row = db.get(
        LotMaterialUse,
        id
    )

    if not row:
        raise HTTPException(
            status_code=404,
            detail="LotMaterialUse not found"
        )

    db.delete(row)
    db.commit()

    return {
        "ok": True
    }
from fastapi import Depends, HTTPException
from fastapi.responses import FileResponse

from sqlalchemy.orm import Session

from docx import Document
from docx.shared import Inches

import tempfile
import qrcode


from docx.shared import Pt

def replace_cell(cell, placeholder, value, size=10):

    if placeholder in cell.text:

        cell.text = ""

        p = cell.paragraphs[0]

        run = p.add_run(str(value))

        run.font.size = Pt(size)

@router.get("/export-docx/{lot_material_use_id}")
def export_docx(
    lot_material_use_id: int,
    qty: int = 4,
    db: Session = Depends(get_db)
):
    print("QTY =", qty)
    row = (
        db.query(
            ProductionLot.lot_no,
            Part.part_no,

            RawBatch.batch_no,
            RawBatch.size_text,
            RawBatch.length_text,

            RawMaterial.type,
            RawMaterial.spec,
        )

        .join(
            LotMaterialUse,
            LotMaterialUse.lot_id == ProductionLot.id
        )

        .join(
            RawBatch,
            RawBatch.id == LotMaterialUse.batch_id
        )

        .join(
            RawMaterial,
            RawMaterial.id == RawBatch.material_id
        )

        .join(
            Part,
            Part.id == ProductionLot.part_id
        )

        .filter(
            LotMaterialUse.id == lot_material_use_id
        )

        .first()
    )

    if not row:
        raise HTTPException(
            status_code=404,
            detail="Record not found"
        )

    # -------------------------
    # Generate QR
    # -------------------------
    qr_text = (
        f"LOT:{row.lot_no}\n"
        f"PART:{row.part_no}\n"
        f"BATCH:{row.batch_no}\n"
        f"TYPE:{row.type}\n"
        f"SPEC:{row.spec}\n"
        f"SIZE:{row.size_text}\n"
        f"LENGTH:{row.length_text}"
    )

    qr = qrcode.QRCode(
        version=1,
        box_size=10,
        border=2
    )

    qr.add_data(qr_text)
    qr.make(fit=True)

    img = qr.make_image(
        fill_color="black",
        back_color="white"
    )

    tmp_qr = tempfile.NamedTemporaryFile(
        delete=False,
        suffix=".png"
    )

    img.save(tmp_qr.name)

    # -------------------------
    # Open Template
    # -------------------------

    if qty == 4: 
        doc = Document(
            "templates/qr_template_mat_4.docx"
        )
    else:
         doc = Document(
            "templates/qr_template.docx"
        )



    mapping = {
        "{{lot}}": row.lot_no or "",
        "{{part}}": row.part_no or "",
        "{{batch}}": row.batch_no or "",
        "{{type}}": row.type or "",
        "{{spec}}": row.spec or "",
        "{{size}}": row.size_text or "",
        "{{length}}": row.length_text or "",
    }

    # -------------------------
    # Replace text in paragraphs
    # -------------------------
    for p in doc.paragraphs:

        for old, new in mapping.items():

            if old in p.text:

                p.text = p.text.replace(
                    old,
                    str(new)
                )

    # -------------------------
    # Replace text in tables
    # -------------------------
    fields = {
        "{{lot}}": (row.lot_no, 22),
        "{{part}}": (row.part_no, 22),
        "{{batch}}": (row.batch_no, 22),
        "{{type}}": (row.type, 22),
        "{{spec}}": (row.spec, 22),
        "{{size}}": (row.size_text, 22),
        "{{length}}": (row.length_text, 22 ),
    }

    for table in doc.tables:
        for row_ in table.rows:
            for cell in row_.cells:

                before = cell.text

                for key, (value, size) in fields.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(
                            key,
                            str(value or "")
                        )

                if before != cell.text:
                    print("REPLACED:", before, "->", cell.text)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(size)

    # -------------------------
    # Insert QR at {{QR}}
    # -------------------------
    for table in doc.tables:
        for row_ in table.rows:
            for cell in row_.cells:

                if "{{QR}}" in cell.text:

                    cell.text = ""

                    p = cell.paragraphs[0]

                    run = p.add_run()

                    run.add_picture(
                        tmp_qr.name,
                        width=Inches(2)
                    )

                    # p.add_run(
                    #     "  " + (row.lot_no or "")
                    # )
    # -------------------------
    # Save
    # -------------------------
    tmp_doc = tempfile.NamedTemporaryFile(
        delete=False,
        suffix=".docx"
    )

    doc.save(tmp_doc.name)

    return FileResponse(
        tmp_doc.name,
        filename=f"{row.lot_no}.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# from docx.shared import Inches
# @router.get("/export-docx/{lot_material_use_id}")
# def export_docx(
#     lot_material_use_id: int,
#     db: Session = Depends(get_db)
# ):

#     row = (
#         db.query(
#             ProductionLot.lot_no,
#             Part.part_no,

#             RawBatch.batch_no,
#             RawBatch.size_text,
#             RawBatch.length_text,

#             RawMaterial.type,
#             RawMaterial.spec,
#         )

#         .join(
#             LotMaterialUse,
#             LotMaterialUse.lot_id == ProductionLot.id
#         )

#         .join(
#             RawBatch,
#             RawBatch.id == LotMaterialUse.batch_id
#         )

#         .join(
#             RawMaterial,
#             RawMaterial.id == RawBatch.material_id
#         )

#         .join(
#             Part,
#             Part.id == ProductionLot.part_id
#         )

#         .filter(
#             LotMaterialUse.id == lot_material_use_id
#         )

#         .first()
#     )

#     if not row:
#         raise HTTPException(
#             status_code=404,
#             detail="Record not found"
#         )

#     doc = Document()

#     doc.add_heading(
#         "Material Traceability",
#         level=1
#     )

#     table = doc.add_table(
#         rows=7,
#         cols=2
#     )

#     table.cell(0, 0).text = "Lot"
#     table.cell(0, 1).text = row.lot_no or ""

#     table.cell(1, 0).text = "Part"
#     table.cell(1, 1).text = row.part_no or ""

#     table.cell(2, 0).text = "Batch No"
#     table.cell(2, 1).text = row.batch_no or ""

#     table.cell(3, 0).text = "Type"
#     table.cell(3, 1).text = row.type or ""

#     table.cell(4, 0).text = "Spec"
#     table.cell(4, 1).text = row.spec or ""

#     table.cell(5, 0).text = "Size"
#     table.cell(5, 1).text = row.size_text or ""

#     table.cell(6, 0).text = "Length"
#     table.cell(6, 1).text = row.length_text or ""

#     tmp = tempfile.NamedTemporaryFile(
#         delete=False,
#         suffix=".docx"
#     )

#     qr_text = (
#         f"LOT:{row.lot_no}\n"
#         f"PART:{row.part_no}\n"
        
#         f"BATCH:{row.batch_no}\n"
#         f"TYPE:{row.type}\n"
#         f"SPEC:{row.spec}\n"
#         f"SIZE:{row.size_text}\n"
        
#         f"LENGTH:{row.length_text}"
#     )

#     qr = qrcode.QRCode(
#         version=1,
#         box_size=10,
#         border=2
#     )

#     qr.add_data(qr_text)
#     qr.make(fit=True)

#     img = qr.make_image(
#         fill_color="black",
#         back_color="white"
#     )

#     tmp_qr = tempfile.NamedTemporaryFile(
#         delete=False,
#         suffix=".png"
#     )

#     img.save(tmp_qr.name)

#     doc.add_paragraph("")
#     doc.add_heading("QR Code", level=2)

#     doc.add_picture(
#         tmp_qr.name,
#         width=Inches(2)
#     )

#     doc.save(tmp.name)

#     return FileResponse(
#         tmp.name,
#         filename=f"{row.lot_no}.docx",
#         media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
#     )


@router.get("/qr/{lot_material_use_id}")
def get_qr(
    lot_material_use_id: int,
    db: Session = Depends(get_db)
):
    import io
    import qrcode

    row = (
        db.query(
            ProductionLot.lot_no,
            Part.part_no,

            RawBatch.batch_no,
            RawBatch.size_text,
            RawBatch.length_text,

            RawMaterial.type,
            RawMaterial.spec,
        )

        .join(
            LotMaterialUse,
            LotMaterialUse.lot_id == ProductionLot.id
        )

        .join(
            RawBatch,
            RawBatch.id == LotMaterialUse.batch_id
        )

        .join(
            RawMaterial,
            RawMaterial.id == RawBatch.material_id
        )

        .join(
            Part,
            Part.id == ProductionLot.part_id
        )

        .filter(
            LotMaterialUse.id == lot_material_use_id
        )

        .first()
    )

    if not row:
        raise HTTPException(
            status_code=404,
            detail="Record not found"
        )

    qr_text = (
        f"LOT:{row.lot_no}\n"
        f"PART:{row.part_no}\n"
        f"BATCH:{row.batch_no}\n"
        f"TYPE:{row.type}\n"
        f"SPEC:{row.spec}\n"
        f"SIZE:{row.size_text}\n"
        f"LENGTH:{row.length_text}"
    )

    img = qrcode.make(qr_text)

    buf = io.BytesIO()

    img.save(
        buf,
        format="PNG"
    )

    buf.seek(0)

    from fastapi.responses import StreamingResponse

    return StreamingResponse(
        buf,
        media_type="image/png"
    )