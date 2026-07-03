# routers/batches.py
from fastapi import APIRouter, Depends, HTTPException, Query
from sqlalchemy.orm import Session
from sqlalchemy.exc import IntegrityError
from sqlalchemy import or_, func, desc
from typing import List, Optional
from pydantic import BaseModel

from database import get_db
from models import RawMaterial, RawBatch,Supplier
from schemas import RawBatchCreate, RawBatchUpdate, RawBatchOut
from utils.code_generator import next_code  # same helper used by customers/materials

router = APIRouter(prefix="/batches", tags=["batches"])

# ---------- Page (offset) ----------
class BatchPage(BaseModel):
    items: List[RawBatchOut]
    total: int
    page: int
    per_page: int
    pages: int

# ---------- Mini for lookup ----------
class BatchMini(BaseModel):
    id: int
    batch_no: Optional[str] = None
    material_id: Optional[int] = None
    class Config:
        from_attributes = True  # Pydantic v2

# ---------- Cursor page (keyset DESC: new -> old) ----------
class BatchCursorPage(BaseModel):
    items: List[RawBatchOut]
    next_cursor: Optional[int] = None   # go older
    prev_cursor: Optional[int] = None   # go newer
    has_more: bool

# ---------- helpers ----------
def _like_escape(term: str) -> str:
    """Escape % and _ for ILIKE and wrap with wildcards."""
    esc = term.replace("\\", "\\\\").replace("%", "\\%").replace("_", "\\_")
    return f"%{esc}%"

# ---------- next batch number (optional for AUTOGEN in UI) ----------
@router.get("/next-no")
def get_next_batch_no(prefix: str = "B", width: int = 5, db: Session = Depends(get_db)):
    return {"next_no": next_code(db, RawBatch, "batch_no", prefix=prefix, width=width)}

# ---------- CREATE (supports AUTO/AUTOGEN/empty batch_no) ----------
@router.post("", response_model=RawBatchOut)
def create_batch(payload: RawBatchCreate, db: Session = Depends(get_db)):
    print("hello")
    # material must exist
    if not db.get(RawMaterial, payload.material_id):
        raise HTTPException(404, "Material not found")

    raw_no = (payload.batch_no or "").strip().upper()
    autogen = raw_no in ("", "AUTO", "AUTOGEN")

    # batch_no = get_next_batch_no( prefix="B", width=5, db = db) if autogen else raw_no
    batch_no = next_code(db, RawBatch, "batch_no", prefix="B", width=5) if autogen else raw_no

    b = RawBatch(
        material_id=payload.material_id,
        batch_no=batch_no,
        supplier_id=payload.supplier_id,

        supplier_batch_no=payload.supplier_batch_no,
        mill_name=payload.mill_name,
        mill_heat_no=payload.mill_heat_no,

        received_at=payload.received_at,
        qty_received=payload.qty_received,
        cert_file=payload.cert_file,
        location=payload.location,

        heat_lot=payload.heat_lot,
        size_text=payload.size_text,
        length_text=payload.length_text,
    )

    for _ in range(3):
        try:
            db.add(b); db.commit(); db.refresh(b)
            return b
        except IntegrityError:
            db.rollback()
            if autogen:
                b.batch_no = next_code(db, RawBatch, "batch_no", prefix="B", width=5)
            else:
                raise HTTPException(409, "Batch number already exists")
    raise HTTPException(500, "Failed to generate unique batch number")

# ---------- LIST (OFFSET) ----------
@router.get("", response_model=BatchPage)
def list_batches(
    q: Optional[str] = Query(None),
    material_id: Optional[int] = Query(None),
    page: int = Query(1, ge=1),
    per_page: int = Query(20, ge=1, le=1000),
    all: Optional[bool] = Query(False),
    db: Session = Depends(get_db),
):
    qry = db.query(RawBatch)

    # Filter by material_id without forcing a join
    if material_id is not None:
        qry = qry.filter(RawBatch.material_id == material_id)

    # If searching across material fields, join only then
    if q and q.strip():
        tokens = q.strip().split()
        qry = qry.join(RawMaterial, RawMaterial.id == RawBatch.material_id)
        for tok in tokens:
            pat = _like_escape(tok)
            # NOTE: func.concat can break on SQLite; switch to || if needed
            qry = qry.filter(or_(
                RawBatch.batch_no.ilike(pat),
                RawBatch.supplier_batch_no.ilike(pat),
                RawBatch.mill_name.ilike(pat),
                RawBatch.mill_heat_no.ilike(pat),
                RawBatch.location.ilike(pat),
                RawMaterial.code.ilike(pat),
                RawMaterial.name.ilike(pat),
                RawMaterial.spec.ilike(pat),
                func.concat(
                    "[", func.coalesce(RawMaterial.code, ""), "] ",
                    func.coalesce(RawMaterial.name, "")
                ).ilike(pat),
            ))

    total = qry.count()

    if all:
        items = qry.order_by(RawBatch.id.desc()).all()
        pages = 1
    else:
        items = (
            qry.order_by(RawBatch.id.desc())
            .offset((page - 1) * per_page)
            .limit(per_page)
            .all()
        )
        pages = (total + per_page - 1) // per_page

    return {
        "items": items,
        "total": total,
        "page": page,
        "per_page": per_page,
        "pages": max(pages, 1),
    }

# ---------- LIST (KEYSET DESC: newest -> oldest) ----------
@router.get("/keyset", response_model=BatchCursorPage)
def list_batches_keyset(
    q: Optional[str] = Query(None, description="Search by batch/material (ILIKE)"),
    material_id: Optional[int] = Query(None, description="Filter by material_id"),
    limit: int = Query(25, ge=1, le=200),
    cursor: Optional[int] = Query(None, description="(DESC) Next page (older): id < cursor"),
    before: Optional[int] = Query(None, description="(DESC) Prev page (newer): id > before"),
    db: Session = Depends(get_db),
):
    
    print("inxxxxxxx...")
    qry = db.query(RawBatch)

    if material_id is not None:
        qry = qry.filter(RawBatch.material_id == material_id)

    if q and q.strip():
        qry = qry.join(RawMaterial, RawMaterial.id == RawBatch.material_id)
        for tok in q.strip().split():
            pat = _like_escape(tok)
            qry = qry.filter(or_(
                RawBatch.batch_no.ilike(pat),
                RawBatch.supplier_batch_no.ilike(pat),
                RawBatch.mill_name.ilike(pat),
                RawBatch.mill_heat_no.ilike(pat),
                RawBatch.location.ilike(pat),
                RawMaterial.code.ilike(pat),
                RawMaterial.name.ilike(pat),
                RawMaterial.spec.ilike(pat),
                func.concat(
                    "[", func.coalesce(RawMaterial.code, ""), "] ",
                    func.coalesce(RawMaterial.name, "")
                ).ilike(pat),
            ))
        # Optional numeric convenience
        if q.strip().isdigit():
            num = int(q.strip())
            qry = qry.filter(or_(RawBatch.id == num, RawBatch.material_id == num))

    # Keyset directions
    going_prev = before is not None and cursor is None
    if going_prev:
        qry = qry.filter(RawBatch.id > before).order_by(RawBatch.id.asc())
        rows = qry.limit(limit + 1).all()
        rows = list(reversed(rows))  # present in DESC order
    else:
        if cursor is not None:
            qry = qry.filter(RawBatch.id < cursor)
        qry = qry.order_by(desc(RawBatch.id))
        rows = qry.limit(limit + 1).all()

    page_rows = rows[:limit]
    has_more = len(rows) > limit
    next_cursor = page_rows[-1].id if page_rows else None
    prev_cursor = page_rows[0].id if page_rows else None

    # Return ORM rows; Pydantic will serialize via from_attributes=True
    return {
        "items": page_rows,
        "next_cursor": next_cursor,
        "prev_cursor": prev_cursor,
        "has_more": has_more,
    }

# ---------- LOOKUP ----------
@router.get("/lookup", response_model=List[BatchMini])
def lookup_batches(ids: str, db: Session = Depends(get_db)):
    try:
        id_list = [int(x) for x in ids.split(",") if x.strip().isdigit()]
    except Exception:
        id_list = []
    if not id_list:
        return []
    rows = db.query(RawBatch).filter(RawBatch.id.in_(id_list)).all()
    return rows

# ---------- GET / UPDATE / DELETE ----------
@router.get("/{batch_id}", response_model=RawBatchOut)
def get_batch(batch_id: int, db: Session = Depends(get_db)):
    b = db.get(RawBatch, batch_id)
    if not b:
        raise HTTPException(404, "Batch not found")
    return b

@router.put("/{batch_id}", response_model=RawBatchOut)
def update_batch(batch_id: int, payload: RawBatchUpdate, db: Session = Depends(get_db)):
    b = db.get(RawBatch, batch_id)
    if not b:
        raise HTTPException(404, "Batch not found")

    data = payload.dict(exclude_unset=True)

    # material change
    if "material_id" in data and data["material_id"] is not None:
        new_mid = int(data["material_id"])
        if not db.get(RawMaterial, new_mid):
            raise HTTPException(404, "Material not found")
        b.material_id = new_mid
        del data["material_id"]

    # qty_received — now independent (no qty_used checks anymore)
    if "qty_received" in data and data["qty_received"] is not None:
        # if using Decimal in model, coerce to string/Decimal in schema instead
        b.qty_received = data["qty_received"]
        del data["qty_received"]

    # other fields
    for k, v in data.items():
        setattr(b, k, v)

    db.commit()
    db.refresh(b)
    return b

@router.delete("/{batch_id}")
def delete_batch(batch_id: int, db: Session = Depends(get_db)):
    b = db.get(RawBatch, batch_id)
    if not b:
        raise HTTPException(404, "Batch not found")
    # If you still track per-part usage via a separate relation, keep that guard.
    if (getattr(b, "uses", []) and len(b.uses) > 0):
        raise HTTPException(400, "Batch already used; cannot delete")
    db.delete(b); db.commit()
    return {"message": "Batch deleted"}

from fastapi.responses import FileResponse
from docx import Document
from docx.shared import Inches
import tempfile
import qrcode

@router.get("/export-docx/{batch_id}")
def export_batch_docx(
    batch_id: int,
    qty: int = 30,
    db: Session = Depends(get_db)
):
    print(qty)   # 4 หรือ 30

    row = (
        db.query(
            RawBatch.batch_no,
            RawBatch.size_text,
            RawBatch.length_text,
            RawBatch.heat_lot,
            RawBatch.date_created,
            RawBatch.heat_po,
            RawBatch.heat_type,

            RawMaterial.type,
            RawMaterial.spec,
            

            Supplier.name.label("supplier_name")
        )

        .join(
            RawMaterial,
            RawMaterial.id == RawBatch.material_id
        )

        .outerjoin(
            Supplier,
            Supplier.id == RawBatch.supplier_id
        )

        .filter(
            RawBatch.id == batch_id
        )

        .first()
    )

    if not row:
        raise HTTPException(
            status_code=404,
            detail="Batch not found"
        )

    # -------------------------
    # Generate QR
    # -------------------------
    print("row", row)
    qr_text = (
        f"BATCH:{row.batch_no}\n"
        f"TYPE:{row.type}\n"
        f"SPEC:{row.spec}\n"
        f"SIZE:{row.size_text}\n"
        f"LENGTH:{row.length_text}\n"
        f"SUPPLIER:{row.supplier_name or ''}"
    )

    qr = qrcode.QRCode(
        version=1,
        box_size=10,
        border=2
    )

    qr.add_data(qr_text)
    qr.make(fit=True)

    img = qr.make_image(
        fill_color="black",
        back_color="white"
    )

    tmp_qr = tempfile.NamedTemporaryFile(
        delete=False,
        suffix=".png"
    )

    img.save(tmp_qr.name)

    # -------------------------
    # Open Template
    # -------------------------
    qr_size = 2
    if qty == 4: 
        doc = Document(
            "templates/qr_template_bat_4.docx"
        )
        qr_size = 1
    elif qty == 30:
        doc = Document(
            "templates/qr_template_bat_30.docx"
        )
        qr_size = 0.8
    elif qty == 80:
         doc = Document(
            "templates/qr_template_bat_80.docx"
        )
         qr_size = 0.4
    print("row.batch_no", row.batch_no)
    mapping = {
        "{{supplier}}" :  row.supplier_name or "",
        "{{batch}}" :  row.batch_no or "",
        "{{mat_po}}": row.batch_no or "",
        "{{type}}": row.type or "",
        "{{spec}}": row.spec or "",
        "{{size}}": row.size_text or "",
        "{{length}}": row.length_text or ""
    }

    print("mapping", mapping)
    # Replace placeholders
    from docx.shared import Pt

    for table in doc.tables:
        for row_ in table.rows:
            for cell in row_.cells:

                for para in cell.paragraphs:
                    for run in para.runs:
                        for old, new in mapping.items():
                            if old in run.text:
                                run.text = run.text.replace(
                                    old,
                                    str(new)
                                )

    # Insert QR
    for table in doc.tables:
        for row_ in table.rows:
            for cell in row_.cells:

                if "{{QR}}" in cell.text:

                    cell.text = ""

                    p = cell.paragraphs[0]

                    run = p.add_run()

                    run.add_picture(
                        tmp_qr.name,
                        width=Inches(qr_size)
                    )

    # -------------------------
    # Save
    # -------------------------
    tmp_doc = tempfile.NamedTemporaryFile(
        delete=False,
        suffix=".docx"
    )

    doc.save(tmp_doc.name)

    return FileResponse(
        tmp_doc.name,
        filename=f"{row.batch_no}.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )