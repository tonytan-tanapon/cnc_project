# routers/v1/icars.py

from fastapi import APIRouter, Depends, HTTPException, Query
from sqlalchemy.orm import Session
from sqlalchemy import or_

from database import get_db
from sqlalchemy.orm import Session, joinedload

from models import (
    ICAR,
    ProductionLot,
    PO
)

router = APIRouter(
    prefix="/icars",
    tags=["ICAR"]
)


from pydantic import BaseModel
from typing import Optional
from datetime import date
from decimal import Decimal


class ICARCreate(BaseModel):
    lot_id: Optional[int] = None
    icar_no: Optional[str] = None

    customer_code: Optional[str] = None
    po_no: Optional[str] = None
    lot_no: Optional[str] = None

    part_no: Optional[str] = None
    rev: Optional[str] = None

    issue_date: Optional[date] = None

    lot_qty: Optional[Decimal] = 0
    defect_qty: Optional[Decimal] = 0
    defect_percent: Optional[Decimal] = 0

    non_conformity: Optional[str] = None

    root_cause: Optional[str] = None

    immediate_corrective_action: Optional[str] = None

    systemic_corrective_action: Optional[str] = None

    preventive_action: Optional[str] = None

    remark: Optional[str] = None

    status: Optional[str] = "open"
    operator_name: Optional[str] = None
    part_name: Optional[str] = None


class ICARUpdate(BaseModel):
    lot_id: Optional[int] = None
    icar_no: Optional[str] = None
    customer_code: Optional[str] = None
    po_no: Optional[str] = None
    lot_no: Optional[str] = None
    part_no: Optional[str] = None
    rev: Optional[str] = None
    issue_date: Optional[date] = None
    lot_qty: Optional[Decimal] = None
    defect_qty: Optional[Decimal] = None
    defect_percent: Optional[Decimal] = None
    non_conformity: Optional[str] = None
    root_cause: Optional[str] = None
    immediate_corrective_action: Optional[str] = None
    systemic_corrective_action: Optional[str] = None
    preventive_action: Optional[str] = None
    remark: Optional[str] = None
    status: Optional[str] = None
    operator_name: Optional[str] = None
    part_name: Optional[str] = None

def to_row(i: ICAR):

    return {
        "id": i.id,
        "lot_id": i.lot_id,
        "icar_no": i.icar_no,
        "customer_code": i.customer_code,
        "po_no": i.po_no,
        "lot_no": i.lot_no,
        "part_no": i.part_no,
        "part_name": i.part_name,
        "rev": i.rev,
        "issue_date": i.issue_date,
        "lot_qty": float(i.lot_qty or 0),
        "defect_qty": float(i.defect_qty or 0),
        "defect_percent": float(i.defect_percent or 0),
        "remark": i.remark,
        "status": i.status,
        "operator_name":
            i.operator_name,
    }

@router.get("/keyset")
def list_keyset(
    q: str = "",
    cursor: int | None = None,
    limit: int = Query(50, le=200),
    db: Session = Depends(get_db),
):

    qry = db.query(ICAR)
    if q:
        like = f"%{q}%"
        qry = qry.filter(
            or_(
                ICAR.icar_no.ilike(like),
                ICAR.customer_code.ilike(like),
                ICAR.po_no.ilike(like),
                ICAR.lot_no.ilike(like),
                ICAR.part_no.ilike(like),
            )
        )

    if cursor:
        qry = qry.filter(ICAR.id < cursor)

    rows = (
        qry.order_by(ICAR.id.desc())
        .limit(limit + 1)
        .all()
    )

    has_more = len(rows) > limit
    rows = rows[:limit]
    next_cursor = None

    if has_more and rows:
        next_cursor = rows[-1].id

    return {
        "items": [to_row(r) for r in rows],
        "next_cursor": next_cursor,
        "has_more": has_more,
    }

@router.get("/{icar_id}")
def get_icar(
    icar_id: int,
    db: Session = Depends(get_db)
):

    row = db.get(ICAR, icar_id)

    if not row:
        raise HTTPException(
            404,
            "ICAR not found"
        )

    return row


@router.post("")
def create_icar(
    payload: ICARCreate,
    db: Session = Depends(get_db)
):

    row = ICAR()

    for k, v in payload.dict().items():
        setattr(row, k, v)

    if not row.issue_date:
        row.issue_date = date.today()

    db.add(row)

    db.commit()

    db.refresh(row)

    return to_row(row)


@router.patch("/{icar_id}")
def update_icar(
    icar_id: int,
    payload: ICARUpdate,
    db: Session = Depends(get_db)
):

    row = db.get(ICAR, icar_id)

    if not row:
        raise HTTPException(
            404,
            "ICAR not found"
        )

    data = payload.dict(
        exclude_unset=True
    )

    for k, v in data.items():
        setattr(row, k, v)

    db.commit()

    db.refresh(row)

    return to_row(row)


@router.delete("/{icar_id}")
def delete_icar(
    icar_id: int,
    db: Session = Depends(get_db)
):

    row = db.get(ICAR, icar_id)

    if not row:
        raise HTTPException(
            404,
            "ICAR not found"
        )

    db.delete(row)

    db.commit()

    return {
        "success": True
    }

@router.get("/lookup/lot/{lot_no}")
def lookup_lot(
    lot_no: str,
    db: Session = Depends(get_db)
):
    lot = (
        db.query(ProductionLot)
        .options(
            joinedload(ProductionLot.po)
            .joinedload(PO.customer),

            joinedload(ProductionLot.part),

            joinedload(
                ProductionLot.part_revision
            )
        )
        .filter(
            ProductionLot.lot_no == lot_no
        )
        .first()
    )

    if not lot:
        raise HTTPException(
            404,
            "Lot not found"
        )

    return {
        "lot_id": lot.id,
        "lot_no": lot.lot_no,

        "customer_code":
            lot.po.customer.code
            if lot.po and lot.po.customer
            else "",

        "po_no":
            lot.po.po_number
            if lot.po
            else "",

        "part_no":
            lot.part.part_no
            if lot.part
            else "",

        "part_name":
            lot.part.name
            if lot.part
            else "",


        "rev":
            lot.part_revision.rev
            if lot.part_revision
            else ""
    }

@router.get("/lots/search")
def search_lots(
    term: str = "",
    db: Session = Depends(get_db)
):

    lots = (
        db.query(ProductionLot)
        .options(
            joinedload(ProductionLot.part),
            joinedload(ProductionLot.part_revision),
            joinedload(ProductionLot.po)
                .joinedload(PO.customer)
        )
        .filter(
            ProductionLot.lot_no.ilike(f"%{term}%")
        )
        .limit(100)
        .all()
    )

    return [

        {
            "label":
                f"{lot.lot_no} | "
                f"{lot.part.part_no if lot.part else ''} | "
                f"{lot.po.customer.code if lot.po and lot.po.customer else ''}",

            "value": lot.lot_no,

            "lot_id": lot.id,

            "customer_code":
                lot.po.customer.code
                if lot.po and lot.po.customer
                else "",

            "po_no":
                lot.po.po_number
                if lot.po
                else "",

            "part_no":
                lot.part.part_no
                if lot.part
                else "",

            "rev":
                lot.part_revision.rev
                if lot.part_revision
                else ""
        }

        for lot in lots
    ]


from docx import Document
from tempfile import NamedTemporaryFile
from fastapi.responses import FileResponse

def replace_text(doc, mapping):

    # =========================
    # PARAGRAPHS
    # =========================

    for p in doc.paragraphs:

        for key, value in mapping.items():

            if key in p.text:

                p.text = p.text.replace(
                    key,
                    str(value)
                )

    # =========================
    # TABLES
    # =========================

    for table in doc.tables:

        for row in table.rows:

            for cell in row.cells:

                for p in cell.paragraphs:

                    for key, value in mapping.items():

                        if key in p.text:

                            p.text = p.text.replace(
                                key,
                                str(value)
                            )

@router.get("/{icar_id}/export-word")
def export_word(
    icar_id: int,
    db: Session = Depends(get_db)
):

    icar = db.get(
        ICAR,
        icar_id
    )
    
    if not icar:
        raise HTTPException(
            404,
            "ICAR not found"
        )

    doc = Document(
        "templates/icar_template.docx"
    )
    
    mapping = {

        "{{icar_no}}":
            icar.icar_no or "",

        "{{issue_date}}":
            icar.issue_date.strftime("%m/%d/%y")
            if icar.issue_date
            else "",

        "{{customer}}":
            icar.customer_code or "",

        "{{po_no}}":
            icar.po_no or "",

        "{{lot_no}}":
            icar.lot_no or "",

        "{{part_no}}":
            icar.part_no or "",

        "{{part_name}}":
            icar.part_name or "",

        "{{rev}}":
            icar.rev or "",

        "{{operator}}":
            icar.operator_name or "",

        "{{lot_qty}}":
            str(
                int(icar.lot_qty or 0)
            ),

        "{{defect_qty}}":
            str(
                int(icar.defect_qty or 0)
            ),

        "{{defect_percent}}":
            f"{float(icar.defect_percent or 0):.2f}%",

        "{{remark}}":
            icar.remark or ""

        
    }

    replace_text(
            doc,
            mapping
        )
    
    print("ICAR")
    print(icar.icar_no)

    print(mapping)

    tmp = NamedTemporaryFile(
        suffix=".docx",
        delete=False
    )

    doc.save(
        tmp.name
    )

    tmp.close()

    return FileResponse(

        tmp.name,

        filename=
            f"ICAR_{icar.icar_no}.docx",

        media_type=
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )